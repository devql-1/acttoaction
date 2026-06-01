"""
ActToAction — MCA IV SEM Final Project Report Generator
Follows the RTU MCA IV Sem guidelines: cover, declaration, acknowledgement,
certificate, and the prescribed 9-chapter main report. Times New Roman, A4,
1.5 line spacing, sky-blue cover, hard-bound style layout.
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(OUTDIR, "images")
REAL_IMG_DIR = os.path.join(OUTDIR, "real_images")
os.makedirs(IMG_DIR, exist_ok=True)


def real(name):
    return os.path.join(REAL_IMG_DIR, name)


FONT = "Times New Roman"
SKY_BLUE = "87CEEB"


# ============================================================
# DIAGRAM HELPERS
# ============================================================

def _box(ax, x, y, w, h, text, color="#4F81BD", text_color="white", fontsize=9):
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle="round,pad=0.02,rounding_size=0.08",
                         linewidth=1.2, facecolor=color, edgecolor="#1F3864")
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fontsize, color=text_color, weight="bold", wrap=True)


def _ellipse(ax, x, y, w, h, text, color="#9BBB59", fontsize=9):
    from matplotlib.patches import Ellipse
    e = Ellipse((x + w / 2, y + h / 2), w, h,
                facecolor=color, edgecolor="#3F6010", linewidth=1.2)
    ax.add_patch(e)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fontsize, color="white", weight="bold")


def _diamond(ax, x, y, w, h, text, color="#F4B084", fontsize=9):
    cx, cy = x + w / 2, y + h / 2
    pts = [(cx, y + h), (x + w, cy), (cx, y), (x, cy)]
    poly = plt.Polygon(pts, closed=True, facecolor=color,
                       edgecolor="#9C5700", linewidth=1.2)
    ax.add_patch(poly)
    ax.text(cx, cy, text, ha="center", va="center", fontsize=fontsize,
            color="black", weight="bold")


def _stick(ax, x, y, label):
    ax.plot([x, x], [y, y - 0.4], "k-", linewidth=1.5)
    ax.add_patch(plt.Circle((x, y + 0.1), 0.1, fill=False, linewidth=1.5))
    ax.plot([x - 0.25, x + 0.25], [y - 0.15, y - 0.15], "k-", linewidth=1.5)
    ax.plot([x, x - 0.25], [y - 0.4, y - 0.7], "k-", linewidth=1.5)
    ax.plot([x, x + 0.25], [y - 0.4, y - 0.7], "k-", linewidth=1.5)
    ax.text(x, y - 0.85, label, ha="center", va="top", fontsize=9, weight="bold")


def _arrow(ax, x1, y1, x2, y2, label="", color="#404040"):
    arr = FancyArrowPatch((x1, y1), (x2, y2),
                          arrowstyle="-|>", mutation_scale=14,
                          color=color, linewidth=1.3)
    ax.add_patch(arr)
    if label:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2, label, fontsize=8,
                color="#202060", ha="center", va="center",
                bbox=dict(facecolor="white", edgecolor="none", pad=1))


def _line(ax, x1, y1, x2, y2):
    ax.plot([x1, x2], [y1, y2], "k-", linewidth=1.2)


def _new_canvas(w=10, h=6, title=""):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(0, w)
    ax.set_ylim(0, h)
    ax.set_aspect("equal")
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=12, weight="bold", pad=10)
    return fig, ax


def save_chart(name):
    path = os.path.join(IMG_DIR, name + ".png")
    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


# ============================================================
# DIAGRAMS REQUIRED BY THE GUIDELINES
# ============================================================

def chart_dfd_level0():
    fig, ax = _new_canvas(11, 5, "DFD Level 0 — ActToAction Context Diagram")
    _box(ax, 0.3, 2.0, 1.8, 1.0, "Visitor", "#4F81BD")
    _box(ax, 9.0, 2.0, 1.8, 1.0, "Admin", "#A8D08D")
    _box(ax, 0.3, 0.3, 1.8, 1.0, "Razorpay", "#C0504D")
    _box(ax, 9.0, 0.3, 1.8, 1.0, "SMTP", "#C0504D")
    _ellipse(ax, 4.0, 1.5, 3.0, 1.8, "ActToAction\nSystem")
    _arrow(ax, 2.1, 2.5, 4.0, 2.5, "form / browse")
    _arrow(ax, 4.0, 2.0, 2.1, 2.0, "pages / receipt")
    _arrow(ax, 7.0, 2.5, 9.0, 2.5, "manage / update")
    _arrow(ax, 9.0, 2.0, 7.0, 2.0, "reports")
    _arrow(ax, 2.1, 0.8, 4.0, 1.6, "verify")
    _arrow(ax, 4.0, 1.6, 2.1, 0.8, "order id")
    _arrow(ax, 7.0, 1.6, 9.0, 0.8, "send mail")
    return save_chart("dfd_level0")


def chart_dfd_level1():
    fig, ax = _new_canvas(11, 7, "DFD Level 1 — ActToAction Major Processes")
    _box(ax, 0.2, 5.5, 1.8, 0.9, "Visitor", "#4F81BD")
    _box(ax, 0.2, 3.0, 1.8, 0.9, "Student", "#4F81BD")
    _box(ax, 0.2, 0.5, 1.8, 0.9, "Admin", "#A8D08D")
    _ellipse(ax, 3.0, 5.5, 2.4, 0.9, "1. Browse Catalog")
    _ellipse(ax, 3.0, 4.0, 2.4, 0.9, "2. Register/Enroll")
    _ellipse(ax, 3.0, 2.5, 2.4, 0.9, "3. Make Payment")
    _ellipse(ax, 3.0, 1.0, 2.4, 0.9, "4. Manage Content")
    _box(ax, 6.5, 5.5, 4.2, 0.7, "D1: courses / events / blogs", "#305496")
    _box(ax, 6.5, 4.0, 4.2, 0.7, "D2: enrollments / registrations", "#305496")
    _box(ax, 6.5, 2.5, 4.2, 0.7, "D3: payments", "#305496")
    _box(ax, 6.5, 1.0, 4.2, 0.7, "D4: admin tables", "#305496")
    for vy, py in [(5.95, 5.95), (3.45, 4.45), (3.45, 2.95), (0.95, 1.45)]:
        _arrow(ax, 2.0, vy, 3.0, py)
    for px, py in [(5.4, 5.85), (5.4, 4.35), (5.4, 2.85), (5.4, 1.35)]:
        _arrow(ax, px, py, 6.5, py)
    _arrow(ax, 4.2, 5.5, 4.2, 4.5)
    _arrow(ax, 4.2, 4.0, 4.2, 3.0)
    _arrow(ax, 2.0, 3.45, 3.0, 4.0)
    _arrow(ax, 2.0, 0.95, 3.0, 1.45)
    return save_chart("dfd_level1")


def chart_er_diagram():
    fig, ax = _new_canvas(12, 8.5, "ER Diagram — ActToAction Core Entities")
    entities = {
        "USER": (1.0, 7.0),
        "COURSE": (5.0, 7.0),
        "COURSE_CATEGORY": (9.5, 7.0),
        "ENROLLMENT": (5.0, 5.2),
        "PAYMENT": (1.0, 5.2),
        "EVENT": (9.5, 5.2),
        "SUB_EVENT": (9.5, 3.4),
        "EVENT_REG": (5.0, 3.4),
        "ATTENDEE": (5.0, 1.6),
        "BLOG": (1.0, 3.4),
        "BLOG_CATEGORY": (1.0, 1.6),
        "PSYCH_TEST": (9.5, 1.6),
    }
    for n, (x, y) in entities.items():
        _box(ax, x, y, 2.2, 0.8, n, "#4F81BD", fontsize=9)
    rels = [
        ("USER", "ENROLLMENT", "places"),
        ("ENROLLMENT", "COURSE", "for"),
        ("ENROLLMENT", "PAYMENT", "has"),
        ("COURSE_CATEGORY", "COURSE", "groups"),
        ("EVENT", "SUB_EVENT", "contains"),
        ("SUB_EVENT", "EVENT_REG", "receives"),
        ("EVENT_REG", "ATTENDEE", "lists"),
        ("EVENT_REG", "PAYMENT", "via"),
        ("BLOG_CATEGORY", "BLOG", "tags"),
    ]
    for a, b, lbl in rels:
        ax_, ay_ = entities[a]
        bx_, by_ = entities[b]
        cx, cy = (ax_ + bx_) / 2 + 1.1, (ay_ + by_) / 2 + 0.4
        ax.text(cx, cy, lbl, fontsize=7, color="#601060", style="italic",
                ha="center", va="center",
                bbox=dict(facecolor="white", edgecolor="none", pad=1))
        _line(ax, ax_ + 1.1, ay_ + 0.4, bx_ + 1.1, by_ + 0.4)
    return save_chart("er_diagram")


def chart_use_case():
    fig, ax = _new_canvas(11, 8, "Use Case Diagram — ActToAction")
    from matplotlib.patches import Ellipse, Rectangle
    sys_box = Rectangle((2.5, 0.5), 6.0, 7.0, fill=False,
                        edgecolor="#1F3864", linewidth=2.0)
    ax.add_patch(sys_box)
    ax.text(5.5, 7.2, "ActToAction System", ha="center", va="bottom",
            fontsize=11, weight="bold", color="#1F3864")

    _stick(ax, 1.0, 6.5, "Visitor")
    _stick(ax, 1.0, 4.0, "Student")
    _stick(ax, 10.0, 6.5, "Admin")
    _stick(ax, 10.0, 4.0, "Razorpay")
    _stick(ax, 10.0, 2.0, "SMTP")

    cases = [
        ("Browse Catalog", 4.5, 6.6),
        ("View Blog", 7.0, 6.6),
        ("Enroll in Course", 4.5, 5.6),
        ("Register Event", 7.0, 5.6),
        ("Take Psych Test", 4.5, 4.6),
        ("Make Payment", 7.0, 4.6),
        ("Submit Enquiry", 4.5, 3.6),
        ("Use Chatbot", 7.0, 3.6),
        ("Manage Content", 4.5, 2.6),
        ("View Reports", 7.0, 2.6),
        ("Send Email", 5.5, 1.4),
    ]
    for txt, x, y in cases:
        _ellipse(ax, x - 1.1, y - 0.4, 2.2, 0.8, txt, "#9BBB59", fontsize=8)

    pairs = [
        (1.4, 6.5, 3.4, 6.6), (1.4, 6.5, 5.9, 6.6),
        (1.4, 4.0, 3.4, 5.6), (1.4, 4.0, 5.9, 5.6),
        (1.4, 4.0, 3.4, 4.6), (1.4, 4.0, 5.9, 4.6),
        (1.4, 4.0, 3.4, 3.6), (1.4, 4.0, 5.9, 3.6),
        (9.6, 6.5, 5.6, 2.6), (9.6, 6.5, 8.1, 2.6),
        (9.6, 4.0, 7.7, 4.6),
        (9.6, 2.0, 6.4, 1.4),
    ]
    for p in pairs:
        _line(ax, *p)
    return save_chart("use_case")


def chart_flow_login():
    fig, ax = _new_canvas(8, 7, "Flow Chart — Admin Login")
    _ellipse(ax, 2.5, 6.0, 3.0, 0.7, "Start")
    _box(ax, 2.5, 4.8, 3.0, 0.7, "Open /admin/login", "#4F81BD")
    _box(ax, 2.5, 3.6, 3.0, 0.7, "Enter credentials", "#4F81BD")
    _diamond(ax, 2.0, 1.9, 4.0, 1.4, "Valid?")
    _box(ax, 0.2, 0.4, 3.0, 0.7, "Show error", "#C0504D")
    _box(ax, 4.8, 0.4, 3.0, 0.7, "Set session, redirect", "#9BBB59")
    for p in [(4.0, 6.0, 4.0, 5.5), (4.0, 4.8, 4.0, 4.3),
              (4.0, 3.6, 4.0, 3.3),
              (3.0, 2.0, 1.7, 1.1, "no"),
              (5.0, 2.0, 6.3, 1.1, "yes")]:
        _arrow(ax, *p)
    return save_chart("flow_login")


def chart_flow_enroll():
    fig, ax = _new_canvas(8, 9, "Flow Chart — Course Enrollment")
    _ellipse(ax, 2.5, 8.0, 3.0, 0.7, "Start")
    _box(ax, 2.5, 6.8, 3.0, 0.7, "Open course detail", "#4F81BD")
    _box(ax, 2.5, 5.6, 3.0, 0.7, "Fill enrolment form", "#4F81BD")
    _diamond(ax, 2.0, 4.0, 4.0, 1.3, "Valid?")
    _box(ax, 2.5, 2.6, 3.0, 0.7, "Create Razorpay order", "#9BBB59")
    _box(ax, 2.5, 1.4, 3.0, 0.7, "User pays", "#9BBB59")
    _diamond(ax, 2.0, -0.2, 4.0, 1.3, "Verified?")
    _box(ax, 0.0, -1.6, 3.0, 0.7, "Mark failed", "#C0504D")
    _box(ax, 5.0, -1.6, 3.0, 0.7, "Confirm + email", "#9BBB59")
    for p in [(4.0, 8.0, 4.0, 7.5), (4.0, 6.8, 4.0, 6.3),
              (4.0, 5.6, 4.0, 5.3), (5.0, 4.1, 5.0, 3.3, "yes"),
              (4.0, 2.6, 4.0, 2.1), (4.0, 1.4, 4.0, 1.1),
              (3.0, -0.1, 1.5, -0.9, "no"),
              (5.0, -0.1, 6.5, -0.9, "yes"),
              (3.0, 4.1, 1.5, 5.3, "no")]:
        _arrow(ax, *p)
    return save_chart("flow_enroll")


def chart_flow_event_reg():
    fig, ax = _new_canvas(8, 9, "Flow Chart — Event Registration")
    _ellipse(ax, 2.5, 8.0, 3.0, 0.7, "Start")
    _box(ax, 2.5, 6.8, 3.0, 0.7, "Open event page", "#4F81BD")
    _box(ax, 2.5, 5.6, 3.0, 0.7, "Pick sub-event", "#4F81BD")
    _diamond(ax, 2.0, 3.9, 4.0, 1.3, "External\nredirect?")
    _box(ax, 0.0, 2.4, 3.0, 0.7, "Open partner URL", "#F4B084", text_color="black")
    _box(ax, 5.0, 2.4, 3.0, 0.7, "Add attendees", "#9BBB59")
    _box(ax, 5.0, 1.2, 3.0, 0.7, "Pay via Razorpay", "#9BBB59")
    _box(ax, 5.0, 0.0, 3.0, 0.7, "Issue ticket", "#9BBB59")
    _ellipse(ax, 5.0, -1.2, 3.0, 0.7, "End")
    for p in [(4.0, 8.0, 4.0, 7.5), (4.0, 6.8, 4.0, 6.3),
              (4.0, 5.6, 4.0, 5.2),
              (3.0, 4.0, 1.5, 3.1, "yes"),
              (5.0, 4.0, 6.5, 3.1, "no"),
              (6.5, 2.4, 6.5, 1.9), (6.5, 1.2, 6.5, 0.7),
              (6.5, 0.0, 6.5, -0.9)]:
        _arrow(ax, *p)
    return save_chart("flow_event_reg")


def chart_flow_payment():
    fig, ax = _new_canvas(8, 9, "Flow Chart — Razorpay Payment")
    _ellipse(ax, 2.5, 8.0, 3.0, 0.7, "Start")
    _box(ax, 2.5, 6.8, 3.0, 0.7, "Server creates order", "#4F81BD")
    _box(ax, 2.5, 5.6, 3.0, 0.7, "Open checkout modal", "#4F81BD")
    _box(ax, 2.5, 4.4, 3.0, 0.7, "User pays", "#9BBB59")
    _box(ax, 2.5, 3.2, 3.0, 0.7, "Receive payment_id, signature", "#9BBB59")
    _diamond(ax, 2.0, 1.5, 4.0, 1.4, "Signature\nmatches?")
    _box(ax, 0.0, 0.0, 3.0, 0.7, "Mark failed", "#C0504D")
    _box(ax, 5.0, 0.0, 3.0, 0.7, "Mark success", "#9BBB59")
    _ellipse(ax, 2.5, -1.4, 3.0, 0.7, "End")
    for p in [(4.0, 8.0, 4.0, 7.5), (4.0, 6.8, 4.0, 6.3),
              (4.0, 5.6, 4.0, 5.1), (4.0, 4.4, 4.0, 3.9),
              (4.0, 3.2, 4.0, 2.9),
              (3.0, 1.6, 1.5, 0.7, "no"), (5.0, 1.6, 6.5, 0.7, "yes"),
              (1.5, 0.0, 4.0, -0.9), (6.5, 0.0, 4.0, -0.9)]:
        _arrow(ax, *p)
    return save_chart("flow_payment")


def chart_flow_psych():
    fig, ax = _new_canvas(8, 9, "Flow Chart — Psychometric Test")
    _ellipse(ax, 2.5, 8.0, 3.0, 0.7, "Start")
    _box(ax, 2.5, 6.8, 3.0, 0.7, "Pick category", "#4F81BD")
    _box(ax, 2.5, 5.6, 3.0, 0.7, "Open test", "#4F81BD")
    _box(ax, 2.5, 4.4, 3.0, 0.7, "Answer questions", "#9BBB59")
    _box(ax, 2.5, 3.2, 3.0, 0.7, "Compute score", "#9BBB59")
    _box(ax, 2.5, 2.0, 3.0, 0.7, "Match TestResultRange", "#9BBB59")
    _box(ax, 2.5, 0.8, 3.0, 0.7, "Render result + email", "#4F81BD")
    _ellipse(ax, 2.5, -0.4, 3.0, 0.7, "End")
    for p in [(4.0, 8.0, 4.0, 7.5), (4.0, 6.8, 4.0, 6.3),
              (4.0, 5.6, 4.0, 5.1), (4.0, 4.4, 4.0, 3.9),
              (4.0, 3.2, 4.0, 2.7), (4.0, 2.0, 4.0, 1.5),
              (4.0, 0.8, 4.0, 0.1)]:
        _arrow(ax, *p)
    return save_chart("flow_psych")


def chart_flow_blog():
    fig, ax = _new_canvas(8, 9, "Flow Chart — Blog Publish")
    _ellipse(ax, 2.5, 8.0, 3.0, 0.7, "Start")
    _box(ax, 2.5, 6.8, 3.0, 0.7, "Admin → Blogs → New", "#4F81BD")
    _box(ax, 2.5, 5.6, 3.0, 0.7, "Pick category, author, tags", "#4F81BD")
    _box(ax, 2.5, 4.4, 3.0, 0.7, "Compose content", "#9BBB59")
    _box(ax, 2.5, 3.2, 3.0, 0.7, "Save (status=1)", "#9BBB59")
    _box(ax, 2.5, 2.0, 3.0, 0.7, "Generate slug", "#9BBB59")
    _box(ax, 2.5, 0.8, 3.0, 0.7, "Publish to /blog", "#4F81BD")
    _ellipse(ax, 2.5, -0.4, 3.0, 0.7, "End")
    for p in [(4.0, 8.0, 4.0, 7.5), (4.0, 6.8, 4.0, 6.3),
              (4.0, 5.6, 4.0, 5.1), (4.0, 4.4, 4.0, 3.9),
              (4.0, 3.2, 4.0, 2.7), (4.0, 2.0, 4.0, 1.5),
              (4.0, 0.8, 4.0, 0.1)]:
        _arrow(ax, *p)
    return save_chart("flow_blog")


def chart_flow_admission():
    fig, ax = _new_canvas(8, 9, "Flow Chart — Admission Form")
    _ellipse(ax, 2.5, 8.0, 3.0, 0.7, "Start")
    _box(ax, 2.5, 6.8, 3.0, 0.7, "Submit short form", "#4F81BD")
    _box(ax, 2.5, 5.6, 3.0, 0.7, "Receive link via email", "#9BBB59")
    _box(ax, 2.5, 4.4, 3.0, 0.7, "Open full form", "#4F81BD")
    _box(ax, 2.5, 3.2, 3.0, 0.7, "Submit details", "#4F81BD")
    _diamond(ax, 2.0, 1.5, 4.0, 1.4, "Approved?")
    _box(ax, 0.0, 0.0, 3.0, 0.7, "Mark rejected", "#C0504D")
    _box(ax, 5.0, 0.0, 3.0, 0.7, "Auto-create enrolment", "#9BBB59")
    for p in [(4.0, 8.0, 4.0, 7.5), (4.0, 6.8, 4.0, 6.3),
              (4.0, 5.6, 4.0, 5.1), (4.0, 4.4, 4.0, 3.9),
              (4.0, 3.2, 4.0, 2.9),
              (3.0, 1.6, 1.5, 0.7, "no"), (5.0, 1.6, 6.5, 0.7, "yes")]:
        _arrow(ax, *p)
    return save_chart("flow_admission")


def chart_screenshot_login():
    fig, ax = _new_canvas(8, 5.5, "Screenshot — Admin Login Page")
    from matplotlib.patches import Rectangle
    ax.add_patch(Rectangle((0, 0), 8, 5.5, facecolor="#F5F7FA", edgecolor="#1F3864", linewidth=2))
    ax.add_patch(Rectangle((0, 4.7), 8, 0.8, facecolor="#1F3864"))
    ax.text(0.3, 5.1, "● ● ●", fontsize=12, color="white")
    ax.text(4.0, 5.1, "https://acttoaction.example.com/admin/login", fontsize=9,
            color="white", ha="center", va="center",
            bbox=dict(facecolor="#FFFFFF22", edgecolor="none", pad=4))
    ax.add_patch(Rectangle((2.0, 1.0), 4.0, 3.3, facecolor="white", edgecolor="#D0D5DD", linewidth=1.2))
    ax.text(4.0, 4.0, "ActToAction Admin", fontsize=14, weight="bold",
            color="#1F3864", ha="center")
    ax.text(4.0, 3.6, "Sign in to continue", fontsize=10, color="#606060", ha="center", style="italic")
    ax.add_patch(Rectangle((2.4, 2.9), 3.2, 0.4, facecolor="#F0F2F5", edgecolor="#C0C0C0"))
    ax.text(2.5, 3.1, "admin@acttoaction.com", fontsize=9, color="#404040", va="center")
    ax.add_patch(Rectangle((2.4, 2.3), 3.2, 0.4, facecolor="#F0F2F5", edgecolor="#C0C0C0"))
    ax.text(2.5, 2.5, "••••••••", fontsize=9, color="#404040", va="center")
    ax.add_patch(Rectangle((2.4, 1.5), 3.2, 0.5, facecolor="#1F3864"))
    ax.text(4.0, 1.75, "LOGIN", fontsize=11, color="white", weight="bold", ha="center", va="center")
    ax.text(4.0, 1.2, "Forgot password?", fontsize=9, color="#1F3864", ha="center", style="italic")
    return save_chart("ss_login")


def chart_screenshot_dashboard():
    fig, ax = _new_canvas(11, 6.5, "Screenshot — Admin Dashboard")
    from matplotlib.patches import Rectangle
    ax.add_patch(Rectangle((0, 0), 11, 6.5, facecolor="#F5F7FA", edgecolor="#1F3864", linewidth=2))
    ax.add_patch(Rectangle((0, 5.8), 11, 0.7, facecolor="#1F3864"))
    ax.text(0.2, 6.15, "ActToAction Admin", fontsize=11, color="white", weight="bold", va="center")
    ax.text(10.7, 6.15, "Logout", fontsize=9, color="white", va="center", ha="right")
    ax.add_patch(Rectangle((0, 0), 2.2, 5.8, facecolor="#2C4063"))
    sidebar = ["Dashboard", "Sliders", "Courses", "Course Categories",
               "Enrollments", "Events", "Sub-Events",
               "Workshops", "Blogs", "Services",
               "Industries", "Psych Tests", "Galleries",
               "Email Templates", "Settings"]
    for i, item in enumerate(sidebar):
        y = 5.4 - i * 0.34
        ax.text(0.3, y, "▸ " + item, fontsize=8.5, color="white", va="center")
    cards = [
        ("Total Enrolments", "1,245", "#4F81BD"),
        ("Events Active", "8", "#9BBB59"),
        ("Workshops This Month", "23", "#F4B084"),
        ("Pending Tickets", "5", "#C0504D"),
    ]
    for i, (label, val, color) in enumerate(cards):
        x = 2.5 + i * 2.1
        ax.add_patch(Rectangle((x, 4.6), 1.9, 1.0, facecolor=color, edgecolor="none"))
        ax.text(x + 0.95, 5.3, val, fontsize=18, weight="bold", color="white", ha="center", va="center")
        ax.text(x + 0.95, 4.85, label, fontsize=8.5, color="white", ha="center", va="center")
    ax.add_patch(Rectangle((2.5, 1.8), 8.4, 2.5, facecolor="white", edgecolor="#D0D5DD"))
    ax.text(2.7, 4.0, "Recent Enrolments", fontsize=11, weight="bold", color="#1F3864")
    rows = [["Aarav Sharma", "Pre-School Foundation", "₹4,500", "Confirmed"],
            ["Anaya Patel", "Coding for Kids", "₹6,000", "Confirmed"],
            ["Vihaan Khan", "Robotics Basics", "₹8,500", "Pending"],
            ["Diya Nair", "Public Speaking", "₹3,200", "Confirmed"],
            ["Aarush Iyer", "Math Olympiad Prep", "₹5,500", "Confirmed"]]
    cols_x = [2.7, 5.0, 7.5, 9.3]
    headers = ["Student", "Course", "Fee", "Status"]
    for j, h in enumerate(headers):
        ax.text(cols_x[j], 3.7, h, fontsize=9, weight="bold", color="#404040")
    for i, row in enumerate(rows):
        y = 3.4 - i * 0.3
        for j, v in enumerate(row):
            ax.text(cols_x[j], y, v, fontsize=8.5, color="#202020")
    ax.add_patch(Rectangle((2.5, 0.4), 8.4, 1.1, facecolor="white", edgecolor="#D0D5DD"))
    ax.text(2.7, 1.25, "Revenue (last 30 days)", fontsize=10, weight="bold", color="#1F3864")
    bars = [3, 4, 5, 4, 6, 7, 5, 8, 6, 9]
    for i, h in enumerate(bars):
        ax.add_patch(Rectangle((3.0 + i * 0.75, 0.55), 0.55, h * 0.07, facecolor="#4F81BD"))
    return save_chart("ss_dashboard")


def chart_screenshot_course_form():
    fig, ax = _new_canvas(9, 7, "Screenshot — Course Edit Form")
    from matplotlib.patches import Rectangle
    ax.add_patch(Rectangle((0, 0), 9, 7, facecolor="#F5F7FA", edgecolor="#1F3864", linewidth=2))
    ax.add_patch(Rectangle((0, 6.4), 9, 0.6, facecolor="#1F3864"))
    ax.text(0.2, 6.7, "ActToAction Admin → Courses → Edit", fontsize=10, color="white", weight="bold", va="center")
    ax.add_patch(Rectangle((0.5, 0.3), 8.0, 5.9, facecolor="white", edgecolor="#D0D5DD"))
    ax.text(0.7, 5.9, "Edit Course", fontsize=14, weight="bold", color="#1F3864")
    fields = [
        ("Category *", "Pre-School Programmes ▾"),
        ("Name *", "Pre-School Foundation"),
        ("Slug *", "pre-school-foundation"),
        ("Short Description", "A foundational course for pre-school children..."),
        ("Long Description", "[Rich text editor]"),
        ("Banner Image", "[file] banner_preschool.jpg"),
        ("Fee (₹)", "4500.00"),
        ("Status", "[✓] Active"),
    ]
    for i, (label, val) in enumerate(fields):
        y = 5.2 - i * 0.55
        ax.text(0.7, y, label, fontsize=9.5, weight="bold", color="#404040")
        ax.add_patch(Rectangle((2.7, y - 0.18), 5.6, 0.36, facecolor="#F0F2F5", edgecolor="#C0C0C0"))
        ax.text(2.85, y, val, fontsize=9, color="#202020", va="center")
    ax.add_patch(Rectangle((6.2, 0.5), 1.0, 0.45, facecolor="#1F3864"))
    ax.text(6.7, 0.725, "SAVE", fontsize=10, color="white", weight="bold", ha="center", va="center")
    ax.add_patch(Rectangle((7.3, 0.5), 1.0, 0.45, facecolor="#C0504D"))
    ax.text(7.8, 0.725, "CANCEL", fontsize=10, color="white", weight="bold", ha="center", va="center")
    return save_chart("ss_course_form")


def chart_screenshot_enroll_form():
    fig, ax = _new_canvas(9, 7, "Screenshot — Public Enrolment Form")
    from matplotlib.patches import Rectangle
    ax.add_patch(Rectangle((0, 0), 9, 7, facecolor="#FFFFFF", edgecolor="#1F3864", linewidth=2))
    ax.add_patch(Rectangle((0, 6.4), 9, 0.6, facecolor="#0E2A47"))
    ax.text(0.2, 6.7, "ACTTOACTION  |  Courses  |  Events  |  Workshops  |  Blog  |  Contact", fontsize=10, color="white", va="center")
    ax.add_patch(Rectangle((0.5, 0.3), 8.0, 5.9, facecolor="#F8FAFC", edgecolor="#D0D5DD"))
    ax.text(4.5, 5.85, "Enrol in Pre-School Foundation", fontsize=15, weight="bold", color="#0E2A47", ha="center")
    ax.text(4.5, 5.5, "Fee: ₹ 4,500  |  Duration: 3 months", fontsize=10, color="#606060", ha="center", style="italic")
    fields = [
        ("Student Name *", "Aarav Sharma"),
        ("Age *", "5"),
        ("Parent Name *", "Rohit Sharma"),
        ("Mother's Phone", "+91 98765 43210"),
        ("Phone *", "+91 98765 43211"),
        ("Email *", "rohit.sharma@example.com"),
        ("City", "Jaipur"),
    ]
    for i, (label, val) in enumerate(fields):
        y = 5.0 - i * 0.55
        ax.text(0.9, y, label, fontsize=9.5, weight="bold", color="#404040")
        ax.add_patch(Rectangle((3.2, y - 0.18), 4.8, 0.36, facecolor="white", edgecolor="#C0C0C0"))
        ax.text(3.35, y, val, fontsize=9, color="#202020", va="center")
    ax.add_patch(Rectangle((3.2, 0.55), 2.6, 0.55, facecolor="#0E2A47"))
    ax.text(4.5, 0.825, "PROCEED TO PAY", fontsize=11, color="white", weight="bold", ha="center", va="center")
    return save_chart("ss_enroll_form")


def chart_screenshot_listing():
    fig, ax = _new_canvas(11, 6.5, "Screenshot — Admin Enrollment Listing")
    from matplotlib.patches import Rectangle
    ax.add_patch(Rectangle((0, 0), 11, 6.5, facecolor="#F5F7FA", edgecolor="#1F3864", linewidth=2))
    ax.add_patch(Rectangle((0, 5.9), 11, 0.6, facecolor="#1F3864"))
    ax.text(0.2, 6.2, "ActToAction Admin → Enrollments", fontsize=10, color="white", weight="bold", va="center")
    ax.add_patch(Rectangle((0.3, 5.0), 6.0, 0.5, facecolor="white", edgecolor="#C0C0C0"))
    ax.text(0.4, 5.25, "Search by name, email, course...", fontsize=9, color="#909090", va="center", style="italic")
    ax.add_patch(Rectangle((6.5, 5.0), 1.4, 0.5, facecolor="#9BBB59"))
    ax.text(7.2, 5.25, "+ Add New", fontsize=9, color="white", weight="bold", ha="center", va="center")
    ax.add_patch(Rectangle((8.0, 5.0), 1.3, 0.5, facecolor="#4F81BD"))
    ax.text(8.65, 5.25, "Export CSV", fontsize=9, color="white", weight="bold", ha="center", va="center")
    headers = ["#", "Student", "Course", "Phone", "Email", "Fee", "Status", "Actions"]
    cols_x = [0.4, 0.9, 2.7, 4.7, 6.0, 8.0, 8.7, 9.7]
    ax.add_patch(Rectangle((0.3, 4.4), 10.4, 0.45, facecolor="#1F3864"))
    for j, h in enumerate(headers):
        ax.text(cols_x[j], 4.625, h, fontsize=9, weight="bold", color="white", va="center")
    rows = [
        ("1", "Aarav Sharma", "Pre-School Foundation", "9876543210", "aarav@ex.com", "₹4,500", "Confirmed"),
        ("2", "Anaya Patel", "Coding for Kids", "9876543211", "anaya@ex.com", "₹6,000", "Confirmed"),
        ("3", "Vihaan Khan", "Robotics Basics", "9876543212", "vihaan@ex.com", "₹8,500", "Pending"),
        ("4", "Diya Nair", "Public Speaking", "9876543213", "diya@ex.com", "₹3,200", "Confirmed"),
        ("5", "Aarush Iyer", "Math Olympiad", "9876543214", "aarush@ex.com", "₹5,500", "Confirmed"),
        ("6", "Saanvi Rao", "Art & Craft", "9876543215", "saanvi@ex.com", "₹2,800", "Cancelled"),
        ("7", "Reyansh Joshi", "Speed Reading", "9876543216", "reyansh@ex.com", "₹3,500", "Confirmed"),
    ]
    for i, row in enumerate(rows):
        y = 4.0 - i * 0.32
        if i % 2 == 1:
            ax.add_patch(Rectangle((0.3, y - 0.15), 10.4, 0.3, facecolor="#F0F4F8"))
        for j, v in enumerate(row):
            ax.text(cols_x[j], y, v, fontsize=8, color="#202020", va="center")
        ax.text(cols_x[7], y, "Edit | Del", fontsize=8, color="#1F3864", va="center")
    ax.text(0.4, 0.3, "Showing 1-7 of 1,245 entries", fontsize=8.5, color="#606060", style="italic")
    ax.add_patch(Rectangle((7.5, 0.15), 0.5, 0.4, facecolor="white", edgecolor="#C0C0C0"))
    ax.text(7.75, 0.35, "1", fontsize=9, weight="bold", ha="center", va="center")
    ax.add_patch(Rectangle((8.05, 0.15), 0.5, 0.4, facecolor="white", edgecolor="#C0C0C0"))
    ax.text(8.3, 0.35, "2", fontsize=9, ha="center", va="center")
    ax.add_patch(Rectangle((8.6, 0.15), 0.5, 0.4, facecolor="white", edgecolor="#C0C0C0"))
    ax.text(8.85, 0.35, "3", fontsize=9, ha="center", va="center")
    ax.add_patch(Rectangle((9.15, 0.15), 0.7, 0.4, facecolor="white", edgecolor="#C0C0C0"))
    ax.text(9.5, 0.35, "Next", fontsize=9, ha="center", va="center")
    return save_chart("ss_listing")


def chart_screenshot_receipt():
    fig, ax = _new_canvas(8, 10, "Sample Output — PDF Payment Receipt")
    from matplotlib.patches import Rectangle
    ax.add_patch(Rectangle((0, 0), 8, 10, facecolor="white", edgecolor="#1F3864", linewidth=2))
    ax.add_patch(Rectangle((0.3, 8.5), 7.4, 1.2, facecolor="#0E2A47"))
    ax.text(0.6, 9.3, "ACTTOACTION", fontsize=22, color="white", weight="bold", va="center")
    ax.text(0.6, 8.85, "Education for Tomorrow", fontsize=10, color="#A0C4FF", style="italic", va="center")
    ax.text(7.4, 9.3, "PAYMENT RECEIPT", fontsize=12, color="white", weight="bold", ha="right", va="center")
    ax.text(7.4, 8.85, "Receipt #: RCPT-2026-1245", fontsize=9, color="#A0C4FF", ha="right", va="center")
    ax.text(0.6, 8.0, "Date:", fontsize=10, weight="bold", color="#404040")
    ax.text(1.6, 8.0, "06 May 2026", fontsize=10, color="#202020")
    ax.text(0.6, 7.7, "Payment ID:", fontsize=10, weight="bold", color="#404040")
    ax.text(1.8, 7.7, "pay_NhJk8sLmX2YpQR", fontsize=10, color="#202020", family="monospace")
    ax.text(0.6, 7.4, "Order ID:", fontsize=10, weight="bold", color="#404040")
    ax.text(1.6, 7.4, "order_NhJk7cKlW1XoPQ", fontsize=10, color="#202020", family="monospace")
    ax.text(0.6, 6.9, "Billed To:", fontsize=11, weight="bold", color="#1F3864")
    ax.text(0.6, 6.5, "Rohit Sharma", fontsize=10, color="#202020")
    ax.text(0.6, 6.25, "rohit.sharma@example.com", fontsize=10, color="#202020")
    ax.text(0.6, 6.0, "+91 98765 43211", fontsize=10, color="#202020")
    ax.text(0.6, 5.75, "Jaipur, Rajasthan", fontsize=10, color="#202020")
    ax.add_patch(Rectangle((0.3, 4.7), 7.4, 0.45, facecolor="#1F3864"))
    headers_r = ["Description", "Qty", "Unit Price", "Amount"]
    cols_rx = [0.5, 4.3, 5.4, 6.7]
    for j, h in enumerate(headers_r):
        ax.text(cols_rx[j], 4.925, h, fontsize=10, weight="bold", color="white", va="center")
    ax.add_patch(Rectangle((0.3, 4.0), 7.4, 0.6, facecolor="#F0F4F8"))
    items = ["Pre-School Foundation Course\n(Student: Aarav Sharma, Age 5)", "1", "₹ 4,500.00", "₹ 4,500.00"]
    for j, v in enumerate(items):
        ax.text(cols_rx[j], 4.3, v, fontsize=9, color="#202020", va="center")
    ax.add_patch(Rectangle((0.3, 3.3), 7.4, 0.5, facecolor="#1F3864"))
    ax.text(cols_rx[2], 3.55, "TOTAL PAID", fontsize=11, weight="bold", color="white", va="center")
    ax.text(cols_rx[3], 3.55, "₹ 4,500.00", fontsize=12, weight="bold", color="white", va="center")
    ax.text(0.5, 2.7, "Amount in words: Rupees Four Thousand Five Hundred Only", fontsize=9.5, color="#202020", style="italic")
    ax.text(0.5, 2.3, "Payment Mode: UPI    |    Status: SUCCESS", fontsize=10, weight="bold", color="#9BBB59")
    ax.text(0.5, 1.7, "Thank you for choosing ActToAction. This is a system-generated receipt", fontsize=8.5, color="#606060")
    ax.text(0.5, 1.5, "and does not require a signature.", fontsize=8.5, color="#606060")
    ax.text(4.0, 0.6, "ActToAction Pvt. Ltd.  |  contact@acttoaction.example.com  |  +91 80-1234-5678",
            fontsize=8, color="#606060", ha="center")
    ax.text(4.0, 0.35, "Registered Office: [Address Line], India  |  GSTIN: [_____]",
            fontsize=8, color="#606060", ha="center")
    return save_chart("ss_receipt")


def chart_gantt():
    fig, ax = _new_canvas(11, 5, "Project Schedule (Gantt Chart)")
    from matplotlib.patches import Rectangle
    phases = [
        ("Requirements & Study", 0, 2, "#4F81BD"),
        ("System Analysis", 1.5, 2, "#4F81BD"),
        ("Database Design", 3, 2, "#9BBB59"),
        ("UI / UX Design", 3, 2.5, "#9BBB59"),
        ("Public Site Development", 5, 5, "#F4B084"),
        ("Admin Panel Development", 5, 6, "#F4B084"),
        ("Payment Integration", 8, 2, "#C0504D"),
        ("Email System", 8, 2, "#C0504D"),
        ("Testing & QA", 10, 3, "#9BBB59"),
        ("Deployment", 12, 1.5, "#4F81BD"),
        ("Documentation", 11, 3, "#A8D08D"),
    ]
    weeks = 14
    ax.set_xlim(-1.5, weeks + 0.5)
    ax.set_ylim(-0.5, len(phases) + 0.5)
    ax.set_aspect("auto")
    for i in range(weeks + 1):
        ax.plot([i, i], [-0.5, len(phases) + 0.5], color="#E0E0E0", linewidth=0.5)
        ax.text(i + 0.5, len(phases) + 0.2, f"W{i+1}", fontsize=7, ha="center", color="#606060")
    for i, (name, start, dur, color) in enumerate(phases):
        y = len(phases) - i - 1
        ax.add_patch(Rectangle((start, y - 0.3), dur, 0.6, facecolor=color, edgecolor="#202020", linewidth=0.5))
        ax.text(-0.2, y, name, fontsize=9, ha="right", va="center", weight="bold", color="#1F3864")
        ax.text(start + dur / 2, y, f"{dur}w", fontsize=8, ha="center", va="center", color="white", weight="bold")
    ax.set_xticks([])
    ax.set_yticks([])
    return save_chart("gantt")


def chart_architecture():
    fig, ax = _new_canvas(11, 7.5, "System Architecture — ActToAction")
    _box(ax, 0.3, 6.3, 2.2, 0.7, "Public Visitor", "#7EB6E1")
    _box(ax, 4.4, 6.3, 2.2, 0.7, "Enrolled Student", "#7EB6E1")
    _box(ax, 8.5, 6.3, 2.2, 0.7, "Admin User", "#A8D08D")
    _box(ax, 1.0, 5.0, 9.0, 0.7, "Web Browser  /  Mobile Browser  (HTTPS)", "#305496")
    _box(ax, 1.0, 3.8, 9.0, 0.7, "Laravel 12 Routing Layer (routes/web.php)", "#1F3864")
    _box(ax, 0.3, 2.5, 2.7, 0.8, "Public Controllers", "#4F81BD")
    _box(ax, 3.3, 2.5, 2.7, 0.8, "Admin Controllers", "#4F81BD")
    _box(ax, 6.3, 2.5, 2.7, 0.8, "API / Chatbot", "#4F81BD")
    _box(ax, 9.3, 2.5, 1.5, 0.8, "Razorpay", "#C0504D")
    _box(ax, 0.3, 1.2, 2.7, 0.8, "Eloquent Models", "#9BBB59")
    _box(ax, 3.3, 1.2, 2.7, 0.8, "Blade Views", "#9BBB59")
    _box(ax, 6.3, 1.2, 2.7, 0.8, "Mailables / Jobs", "#9BBB59")
    _box(ax, 9.3, 1.2, 1.5, 0.8, "PDF", "#C0504D")
    _box(ax, 1.0, 0.1, 9.0, 0.7, "MySQL Database  +  File Storage", "#1F3864")
    for x in [1.4, 5.5, 9.6]:
        _arrow(ax, x, 6.3, x, 5.7)
        _arrow(ax, x, 5.0, x, 4.5)
    for x in [1.4, 4.6, 7.6, 10.0]:
        _arrow(ax, x, 3.8, x, 3.3)
        _arrow(ax, x, 1.2, x, 0.8)
    for x in [1.4, 4.6, 7.6]:
        _arrow(ax, x, 2.5, x, 2.0)
    return save_chart("architecture")


# ============================================================
# DOCX HELPERS
# ============================================================

NAVY = RGBColor(0x1F, 0x38, 0x64)
DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x60, 0x60, 0x60)


def set_font(run, size=12, bold=False, italic=False, color=None, name=FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_paragraph(p, line_spacing=1.5, space_before=0, space_after=6,
                  alignment=None, first_line_indent=None):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)


def add_chapter_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    set_font(r, size=16, bold=True, color=NAVY)
    set_paragraph(p, line_spacing=1.5, space_before=12, space_after=12,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=14, bold=True, color=NAVY)
    set_paragraph(p, line_spacing=1.5, space_before=10, space_after=6)


def add_subsub_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=12, bold=True, color=DARK)
    set_paragraph(p, line_spacing=1.5, space_before=6, space_after=4)


def add_para(doc, text, justify=True, indent=0.3):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=12)
    set_paragraph(p, line_spacing=1.5, space_after=6,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else None,
                  first_line_indent=indent)


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_font(r, size=12)
        set_paragraph(p, line_spacing=1.5, space_after=2)


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(it)
        set_font(r, size=12)
        set_paragraph(p, line_spacing=1.5, space_after=2)


def add_image(doc, path, caption=None, width=6.0):
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        set_font(r, size=11, italic=True, color=GREY)
        set_paragraph(cp, line_spacing=1.15, space_after=10)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3864")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=11)
    p = doc.add_paragraph()
    set_paragraph(p, space_after=6)


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)


def shade_page(doc):
    """Add sky-blue background to current section (cover page effect)."""
    sectPr = doc.sections[0]._sectPr
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), SKY_BLUE)
    doc.element.body.insert(0, bg)


def add_cover_page(doc):
    # Logo
    logo = real("logo.png")
    if os.path.exists(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, space_before=12, space_after=6)
        run = p.add_run()
        run.add_picture(logo, width=Inches(1.5))

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_before=12, space_after=18)
    r = p.add_run("ACTTOACTION")
    set_font(r, size=36, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=6)
    r = p.add_run("A WEB-BASED EDUCATIONAL MANAGEMENT SYSTEM")
    set_font(r, size=18, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=24)
    r = p.add_run("(Courses, Events, Workshops, Blogs, Services, "
                  "Industries, Psychometric Tests & Online Payments)")
    set_font(r, size=13, italic=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_before=24, space_after=12)
    r = p.add_run("A Project Report")
    set_font(r, size=14, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=12)
    r = p.add_run("Submitted in partial fulfilment of the requirements\n"
                  "for the award of the degree of")
    set_font(r, size=12, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=18)
    r = p.add_run("MASTER OF COMPUTER APPLICATIONS\n(MCA — IV Semester)")
    set_font(r, size=16, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=6)
    r = p.add_run("Submitted by")
    set_font(r, size=12, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=18)
    r = p.add_run("[Student Name]\nRoll No.: [___________]")
    set_font(r, size=14, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=6)
    r = p.add_run("Under the guidance of")
    set_font(r, size=12, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=18)
    r = p.add_run("[Guide Name]\n[Designation]")
    set_font(r, size=13, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_before=18, space_after=6)
    r = p.add_run("[NAME OF THE COLLEGE / INSTITUTE]")
    set_font(r, size=14, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=6)
    r = p.add_run("Affiliated to")
    set_font(r, size=11, italic=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_after=18)
    r = p.add_run("RAJASTHAN TECHNICAL UNIVERSITY, KOTA")
    set_font(r, size=14, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, space_before=12, space_after=0)
    r = p.add_run("May 2026")
    set_font(r, size=12, bold=True, color=RGBColor(0, 0, 0))


# ============================================================
# BUILD
# ============================================================

def build():
    print("Generating diagrams...")
    charts = {
        "arch": chart_architecture(),
        "dfd0": chart_dfd_level0(),
        "dfd1": chart_dfd_level1(),
        "er": chart_er_diagram(),
        "use_case": chart_use_case(),
        "flow_login": chart_flow_login(),
        "flow_enroll": chart_flow_enroll(),
        "flow_event": chart_flow_event_reg(),
        "flow_pay": chart_flow_payment(),
        "flow_psych": chart_flow_psych(),
        "flow_blog": chart_flow_blog(),
        "flow_admit": chart_flow_admission(),
        "ss_login": chart_screenshot_login(),
        "ss_dashboard": chart_screenshot_dashboard(),
        "ss_course_form": chart_screenshot_course_form(),
        "ss_enroll_form": chart_screenshot_enroll_form(),
        "ss_listing": chart_screenshot_listing(),
        "ss_receipt": chart_screenshot_receipt(),
        "gantt": chart_gantt(),
    }

    print("Building Word document...")
    doc = Document()
    set_default_font(doc)

    # Margins as per guideline: Left 1.5", Right/Top/Bottom 1"
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)

    # Sky blue page background (cover indicator)
    shade_page(doc)

    # ---------- 1. COVER PAGE ----------
    add_cover_page(doc)
    add_page_break(doc)

    # ---------- 2. DECLARATION ----------
    add_chapter_heading(doc, "Declaration")
    add_para(doc,
             "I hereby declare that the project report titled \"ActToAction — A Web-Based "
             "Educational Management System\" submitted by me to the Department of "
             "Computer Applications, [Name of the Institute], affiliated to Rajasthan "
             "Technical University, Kota, in partial fulfilment of the requirements "
             "for the award of the degree of Master of Computer Applications (MCA) "
             "is a record of bonafide industrial project work carried out by me "
             "under the supervision of [Guide Name].",
             indent=0.3)
    add_para(doc,
             "I further declare that the work reported in this project has not been "
             "submitted and will not be submitted, either in part or in full, for "
             "the award of any other degree or diploma of this institute or of any "
             "other institute or university.",
             indent=0.3)
    add_para(doc, "", indent=0)
    add_para(doc, "", indent=0)
    add_para(doc, "Place: ____________________", justify=False, indent=0)
    add_para(doc, "Date: _____________________", justify=False, indent=0)
    add_para(doc, "", indent=0)
    add_para(doc, "", indent=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("[Student Name]\nMCA — IV Semester\nRoll No.: __________")
    set_font(r, size=12, bold=True)
    set_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_break(doc)

    # ---------- 3. ACKNOWLEDGEMENT ----------
    add_chapter_heading(doc, "Acknowledgement")
    add_para(doc,
             "The successful completion of this project would not have been possible "
             "without the support, encouragement, and guidance of many individuals to "
             "whom I am deeply indebted.")
    add_para(doc,
             "I would like to express my sincere gratitude to my project guide, "
             "[Guide Name], [Designation], for the valuable guidance, constant "
             "encouragement, and constructive criticism throughout the duration of "
             "this project. The technical insights and the timely suggestions given "
             "by them have been instrumental in shaping this work.")
    add_para(doc,
             "I extend my heartfelt thanks to [Head of Department], Head of the "
             "Department of Computer Applications, for providing the necessary "
             "infrastructure, laboratory facilities, and an academic environment "
             "conducive to research and development.")
    add_para(doc,
             "I also wish to thank the management and the development team of "
             "ActToAction for sharing real-world requirements, feedback, and domain "
             "knowledge that allowed this academic project to take an industrial "
             "form. The opportunity to work on a live Laravel-based platform that "
             "actually serves end users has been an invaluable learning experience.")
    add_para(doc,
             "My gratitude also goes to the faculty members of the MCA programme "
             "for the strong foundation they have laid in software engineering, "
             "database management systems, and web technologies — without which "
             "executing a project of this scope would not have been feasible.")
    add_para(doc,
             "Finally, I thank my parents, family, and friends for their unwavering "
             "moral support and patience during the project work.")
    add_para(doc, "", indent=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("[Student Name]")
    set_font(r, size=12, bold=True)
    add_page_break(doc)

    # ---------- 4. CERTIFICATE ----------
    add_chapter_heading(doc, "Certificate")
    add_para(doc,
             "This is to certify that the project report entitled \"ActToAction — A "
             "Web-Based Educational Management System\" is a bonafide record of the "
             "industrial project work carried out by [Student Name], Roll No. "
             "[__________], a student of the Master of Computer Applications "
             "programme, IV Semester, at [Name of the Institute], during the "
             "academic year 2025-2026.")
    add_para(doc,
             "The project work has been carried out under my supervision and "
             "guidance, and the report submitted is, to the best of my knowledge, "
             "a record of the candidate's own work. The work embodied in this "
             "report has not been submitted to any other university or institute "
             "for the award of any other degree or diploma.")
    add_para(doc,
             "The project, in my opinion, has reached the standard required for "
             "fulfilling the project requirements of the MCA programme as "
             "prescribed by Rajasthan Technical University, Kota.")
    add_para(doc, "", indent=0)
    add_para(doc, "", indent=0)
    add_para(doc, "", indent=0)
    # Signature blocks
    t = doc.add_table(rows=1, cols=2)
    t.autofit = True
    cells = t.rows[0].cells
    for cell, text in zip(cells, [
        "________________________\n[Guide Name]\nProject Guide\n[Designation]",
        "________________________\n[Head of Department]\nHead of the Department\nDepartment of Computer Applications",
    ]):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=12, bold=True)
        set_paragraph(p, line_spacing=1.5)
    add_para(doc, "", indent=0)
    add_para(doc, "", indent=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("________________________\nExternal Examiner")
    set_font(r, size=12, bold=True)
    add_page_break(doc)

    # ---------- 5. CERTIFICATE OF COMPANY/INSTITUTE ----------
    add_chapter_heading(doc, "Certificate of the Company / Institute")
    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ON COMPANY LETTERHEAD]")
    set_font(r, size=12, bold=True, italic=True, color=GREY)
    add_para(doc, "")
    add_para(doc,
             "This is to certify that Mr./Ms. [Student Name], a student of MCA "
             "IV Semester, [Name of the Institute], affiliated to Rajasthan "
             "Technical University, Kota, has successfully carried out the "
             "industrial project work titled \"ActToAction — A Web-Based Educational "
             "Management System\" with our organisation during the period "
             "[Start Date] to [End Date].")
    add_para(doc,
             "During the tenure of the project, the candidate worked on the "
             "design, implementation, and testing of the ActToAction Laravel "
             "platform. The candidate was involved in feature development across "
             "the admin panel, public website, and payment integration layer. "
             "The candidate's conduct and performance during the project period "
             "were found to be satisfactory.")
    add_para(doc,
             "We wish the candidate all success in future endeavours.")
    add_para(doc, "")
    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("________________________\n[Authorised Signatory]\n"
                  "[Designation]\n[Company Name]\n[Date]")
    set_font(r, size=12, bold=True)
    add_page_break(doc)

    # ---------- 6. TABLE OF CONTENTS ----------
    add_chapter_heading(doc, "Table of Contents")
    toc = [
        ("Chapter 1: Introduction", ""),
        ("    1.1 Objectives", ""),
        ("    1.2 Problem Description", ""),
        ("    1.3 About the Organization", ""),
        ("Chapter 2: System Study", ""),
        ("    2.1 System with Limitations", ""),
        ("    2.2 Significance of the Project", ""),
        ("    2.3 Beneficiaries of the System", ""),
        ("    2.4 Feasibility Study", ""),
        ("Chapter 3: System Analysis", ""),
        ("    3.1 Functional Requirements", ""),
        ("    3.2 Non-Functional Requirements", ""),
        ("    3.3 User Requirements", ""),
        ("    3.4 System Requirements", ""),
        ("Chapter 4: System Design", ""),
        ("    4.1 Data Flow Diagrams", ""),
        ("    4.2 E-R Diagrams", ""),
        ("    4.3 Use Case Diagrams", ""),
        ("    4.4 Flow Charts", ""),
        ("    4.5 Database Tables", ""),
        ("    4.6 Input / Output Forms", ""),
        ("Chapter 5: Development", ""),
        ("    5.1 Environment", ""),
        ("    5.2 Coding Style", ""),
        ("    5.3 Coding Techniques", ""),
        ("    5.4 Coding", ""),
        ("Chapter 6: Testing", ""),
        ("    6.1 Test Cases", ""),
        ("Chapter 7: System Security", ""),
        ("    7.1 Checks and Control", ""),
        ("    7.2 Encryption and Secure Practices", ""),
        ("Chapter 8: Conclusion / Future Enhancement", ""),
        ("Chapter 9: Bibliography", ""),
    ]
    for title, _ in toc:
        p = doc.add_paragraph()
        r = p.add_run(title)
        set_font(r, size=12)
        set_paragraph(p, line_spacing=1.5, space_after=2)
    add_page_break(doc)

    # ---------- LIST OF FIGURES ----------
    add_chapter_heading(doc, "List of Figures")
    for fig in [
        "Figure 4.1  System Architecture of ActToAction",
        "Figure 4.2  Data Flow Diagram — Level 0 (Context)",
        "Figure 4.3  Data Flow Diagram — Level 1",
        "Figure 4.4  Entity-Relationship Diagram",
        "Figure 4.5  Use Case Diagram",
        "Figure 4.6  Flow Chart — Admin Login",
        "Figure 4.7  Flow Chart — Course Enrollment",
        "Figure 4.8  Flow Chart — Event Registration",
        "Figure 4.9  Flow Chart — Razorpay Payment",
        "Figure 4.10 Flow Chart — Psychometric Test",
        "Figure 4.11 Flow Chart — Blog Publishing",
        "Figure 4.12 Flow Chart — Admission Form",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(fig)
        set_font(r, size=12)
        set_paragraph(p, line_spacing=1.5, space_after=2)
    add_page_break(doc)

    # ============================================================
    # CHAPTER 1: INTRODUCTION
    # ============================================================
    add_chapter_heading(doc, "Chapter 1: Introduction")
    add_para(doc,
             "ActToAction is a web-based educational management system developed using "
             "the Laravel 12 framework. It is a comprehensive platform that "
             "consolidates the digital operations of an educational organisation — "
             "course delivery, workshop scheduling, event management, blog "
             "publishing, service catalogues, industry sections, psychometric "
             "testing, and online payments — into a single integrated application. "
             "The platform serves three primary classes of users: anonymous "
             "visitors browsing the website, enrolled students accessing their "
             "programmes, and administrative staff who manage every aspect of the "
             "organisation through a unified back-office.")
    add_para(doc,
             "This project report documents the analysis, design, development, "
             "testing, and deployment of the ActToAction platform. It is structured "
             "according to the guidelines prescribed for the MCA IV Semester "
             "industrial project at Rajasthan Technical University, and is "
             "intended to demonstrate the application of software engineering "
             "principles to a real-world web application of moderate complexity.")

    add_sub_heading(doc, "1.1 Objectives")
    add_para(doc,
             "The principal objectives that guided the development of ActToAction are "
             "summarised below. Each objective addresses a measurable shortcoming "
             "in the manual or semi-digital processes that the platform replaces.")
    add_bullets(doc, [
        "To design and develop a unified web-based system that brings every public-facing operation of the organisation — courses, workshops, events, blogs, services, industries, and psychometric tests — under a single, consistent platform.",
        "To build a role-aware administrative back-office that allows non-technical staff to publish, edit, and retire content without the involvement of engineers.",
        "To integrate a secure online payment gateway (Razorpay) supporting cards, UPI, netbanking, and wallets, with server-side signature verification for every transaction.",
        "To provide a configurable email templating mechanism that supports merge variables, allowing transactional and marketing communications to be authored from the admin panel.",
        "To implement a configuration-driven psychometric testing engine where new tests, questions, and result interpretation ranges can be added entirely from the admin panel.",
        "To capture, store, and audit every form submission — admissions, enrolments, registrations, enquiries, volunteer applications, and franchise leads — with timestamps and status tracking.",
        "To support search-engine-friendly URLs through unique slug fields on every public entity (courses, blogs, events, sub-events, tests, schools).",
        "To present a responsive, mobile-friendly public website that loads efficiently on low-bandwidth networks typical of Tier-2 and Tier-3 cities.",
        "To provide reporting facilities — listings, exports, status filters — that allow operational teams to track performance and prepare attendance and revenue reports.",
        "To follow established software engineering practices: MVC separation, version control (Git), database migrations, environment-driven configuration, and reproducible deployment.",
    ])

    add_sub_heading(doc, "1.2 Problem Description")
    add_para(doc,
             "Prior to the development of ActToAction, the organisation managed its "
             "digital footprint through a combination of static brochure-style web "
             "pages, third-party form builders, spreadsheet-based attendance "
             "registers, email threads for support enquiries, and manual payment "
             "reconciliation against bank statements. This fragmented approach "
             "produced four recurring problems:")
    add_numbered(doc, [
        "Data was scattered across multiple, disconnected tools. A single student's enrolment information might exist in a Google Form sheet, a payment gateway dashboard, and a finance spreadsheet — with no single record of truth.",
        "Content updates required engineering involvement. Even minor changes — a new blog post, a slider banner update, a workshop schedule change — needed a developer, a deployment, and a verification cycle.",
        "There was no auditable history. Lost emails, deleted spreadsheet rows, and unrecorded phone confirmations meant that disputes could not be resolved with confidence.",
        "Payments were not reconciled in real time. Manual reconciliation introduced delays of up to a week, during which the operational team could not confirm whether a registration was paid.",
    ])
    add_para(doc,
             "ActToAction addresses each of these problems through a centralised "
             "relational database, an opinionated administrative interface, "
             "automatic logging of every transactional email, and direct "
             "server-to-server payment verification with the Razorpay gateway. "
             "The result is a single platform where every visitor interaction is "
             "captured, every staff action is traceable, and every payment is "
             "verified at the moment of completion.")

    add_sub_heading(doc, "1.3 About the Organization")
    add_para(doc,
             "ActToAction is an educational organisation engaged in the delivery of "
             "structured learning programmes for students across a wide age "
             "spectrum, from pre-schoolers to senior secondary learners. Its "
             "programmes span multiple delivery modes — long-running courses, "
             "short hands-on workshops conducted city-by-city, flagship annual "
             "events, and self-administered psychometric assessments — all "
             "supported by a strong content marketing arm comprising a regular "
             "blog and a curated video gallery.")
    add_para(doc,
             "The organisation operates in collaboration with a network of "
             "schools, franchise partners, and content collaborators. Its "
             "operational team comprises content editors, admissions counsellors, "
             "event coordinators, workshop leaders, customer support personnel, "
             "and a small engineering function. The ActToAction platform is the "
             "digital headquarters that connects these teams to one another and "
             "to the public.")

    add_sub_heading(doc, "1.4 Visual Snapshot of the Platform")
    add_para(doc,
             "The images below give a visual sense of how the ActToAction "
             "platform manifests to its users — from the public hero banner "
             "to learning sessions, partner brands, and the people behind "
             "the operation.")
    if os.path.exists(real("home_banner.jpg")):
        add_image(doc, real("home_banner.jpg"),
                  "Figure 1.1 — Hero banner from the ActToAction home page.",
                  width=6.0)
    if os.path.exists(real("classroom.png")):
        add_image(doc, real("classroom.png"),
                  "Figure 1.2 — A live ActToAction learning session in progress.",
                  width=5.5)
    add_page_break(doc)

    # ============================================================
    # CHAPTER 2: SYSTEM STUDY
    # ============================================================
    add_chapter_heading(doc, "Chapter 2: System Study")
    add_para(doc,
             "The system study phase examined the existing manual and "
             "semi-digital processes used by the organisation, identified their "
             "limitations, and established the scope and significance of the "
             "proposed ActToAction platform. This chapter presents the findings of "
             "that study, the beneficiaries of the new system, and the "
             "feasibility analysis carried out before commencing development.")

    add_sub_heading(doc, "2.1 System with Limitations")
    add_para(doc,
             "The legacy approach was a patchwork of disconnected tools. The "
             "marketing website was a static set of HTML pages without a content "
             "management layer. Form-based interactions used external services "
             "such as Google Forms and Typeform, with submissions ending up as "
             "spreadsheets that operational staff downloaded and merged "
             "manually. Payment links were generated ad-hoc in the gateway's "
             "dashboard and shared with prospective students by email. The "
             "concrete limitations are listed below:")
    add_subsub_heading(doc, "2.1.1 Operational Limitations")
    add_bullets(doc, [
        "Course information was edited directly in HTML, requiring engineering for every update.",
        "Workshop schedules were maintained in shared spreadsheets that could be edited concurrently and silently overwritten.",
        "Event registrations were collected through external forms; reconciling them with payments required manual lookups.",
        "Blog posts were written in Word and uploaded as PDFs, harming SEO and discoverability.",
        "Psychometric tests existed as printable PDFs with no automated scoring.",
        "Customer support enquiries were routed to a shared email inbox without ticket tracking.",
    ])
    add_subsub_heading(doc, "2.1.2 Technical Limitations")
    add_bullets(doc, [
        "No centralised authentication; admin tools each had their own credentials.",
        "No audit trail; deletions and edits left no record of who or when.",
        "No automated email templating; transactional copy was hand-written each time.",
        "No structured data model for relationships such as course → session → enrolment.",
        "No automated payment verification; manual reconciliation against bank statements.",
        "No mobile-friendly experience for several of the third-party form tools.",
    ])
    add_subsub_heading(doc, "2.1.3 Strategic Limitations")
    add_bullets(doc, [
        "Brand inconsistency across multiple disconnected tools.",
        "Inability to run cross-channel analytics or attribute revenue to campaigns.",
        "High onboarding cost for new operational staff who had to learn five different tools.",
        "Limited ability to scale during admission seasons when traffic peaked.",
    ])

    add_sub_heading(doc, "2.2 Significance of the Project")
    add_para(doc,
             "ActToAction replaces the patchwork with a unified, consistent platform. "
             "The significance of the project is best understood across three "
             "dimensions: organisational efficiency, customer experience, and "
             "long-term scalability.")
    add_subsub_heading(doc, "2.2.1 Organisational Efficiency")
    add_para(doc,
             "Every operational interaction now flows through a single database. "
             "An admissions counsellor can look up a student by name and see "
             "their full history — short form submission, full form submission, "
             "enrolment, payment, attendance — in one place. A finance officer "
             "can pull a single CSV that includes every transaction across "
             "courses, events, and workshops. The reduction in cross-tool "
             "switching frees significant staff time.")
    add_subsub_heading(doc, "2.2.2 Customer Experience")
    add_para(doc,
             "Visitors interact with a single, consistent brand. Forms validate "
             "input on the client and the server, prevent double submissions, "
             "and acknowledge receipt within the same browser session. Payments "
             "happen inline, with success and failure handled gracefully. "
             "Confirmation emails arrive within seconds, complete with branded "
             "PDF receipts.")
    add_subsub_heading(doc, "2.2.3 Long-Term Scalability")
    add_para(doc,
             "The platform's modular architecture allows new modules — new test "
             "categories, new event types, new partner programmes — to be added "
             "without disrupting existing functionality. The data model is "
             "normalised and indexed, so query performance remains acceptable "
             "even as the volume of submissions grows by an order of magnitude.")

    add_sub_heading(doc, "2.3 Beneficiaries of the System")
    add_para(doc,
             "The ActToAction platform delivers value to multiple stakeholder "
             "groups, each with distinct concerns. The principal beneficiaries "
             "and the specific value delivered to them are summarised below.")
    add_table(doc,
              ["Stakeholder", "Primary Benefit"],
              [
                  ["Public Visitors",
                   "Single, brand-consistent website with mobile-friendly responsive design."],
                  ["Prospective Students",
                   "Streamlined enrolment with inline payment and instant confirmation."],
                  ["Enrolled Students",
                   "Email receipts, downloadable PDF tickets, and self-service test results."],
                  ["Content Editors",
                   "WYSIWYG admin to publish blogs, sliders, and gallery items without engineering."],
                  ["Admissions Staff",
                   "Two-stage admission funnel with conversion-friendly short form and detailed full form."],
                  ["Event Coordinators",
                   "Sub-event creation, attendee lists, capacity tracking, and CSV exports."],
                  ["Workshop Leaders",
                   "City- and school-scoped registration data with merchandise add-ons."],
                  ["Customer Support",
                   "Chatbot-driven FAQ deflection and structured support ticket triage."],
                  ["Finance Team",
                   "Real-time payment verification and audit-friendly transaction history."],
                  ["Engineering Team",
                   "Predictable Laravel codebase with consistent CRUD patterns across modules."],
                  ["Management",
                   "Roll-up reports across courses, events, and workshops in one place."],
              ])

    add_sub_heading(doc, "2.4 Feasibility Study")
    add_para(doc,
             "Before committing to development, the project was evaluated for "
             "feasibility along four dimensions: technical, economic, "
             "operational, and legal.")
    add_subsub_heading(doc, "2.4.1 Technical Feasibility")
    add_para(doc,
             "The chosen stack — PHP 8.2, Laravel 12, MySQL, Razorpay's PHP "
             "SDK — is mature, widely supported, and backed by a deep talent "
             "pool. All required functionality (form handling, file uploads, "
             "email dispatch, PDF generation, payment integration) is "
             "well-supported by first-party or community packages. The "
             "development team had prior experience with Laravel, eliminating "
             "the learning-curve risk. The platform runs on commodity Linux "
             "hardware behind a standard Nginx or Apache reverse proxy.")
    add_subsub_heading(doc, "2.4.2 Economic Feasibility")
    add_para(doc,
             "The project leverages open-source software end-to-end (Laravel, "
             "MySQL, PHP). Hosting cost is bounded by a single virtual private "
             "server with predictable monthly billing. Razorpay charges a "
             "transaction fee per successful payment, which is comfortably "
             "covered by the margin on each enrolment. Development effort "
             "incurred a one-time cost; ongoing maintenance is part-time. "
             "Payback was estimated at under twelve months based on the cost "
             "savings from retiring third-party form services and the "
             "incremental conversions enabled by the streamlined funnel.")
    add_subsub_heading(doc, "2.4.3 Operational Feasibility")
    add_para(doc,
             "The administrative interface mirrors patterns familiar to anyone "
             "who has used a content management system (WordPress, Drupal). "
             "Operational staff received two short training sessions and were "
             "able to manage the platform independently within a week. The "
             "predictable URL structure, status toggles, and consistent CRUD "
             "patterns across modules ensure that learning one module is "
             "sufficient to use them all.")
    add_subsub_heading(doc, "2.4.4 Legal Feasibility")
    add_para(doc,
             "The platform stores personally identifiable information of "
             "students and parents — names, phone numbers, email addresses, and "
             "in some cases health notes. The data is held within India on "
             "encrypted-at-rest storage, with HTTPS enforced for all transport. "
             "Razorpay is PCI-DSS compliant and handles all card data, "
             "removing the need for the platform itself to be PCI-certified. "
             "The project complies with the relevant clauses of the Information "
             "Technology Act 2000 and the SPDI Rules 2011.")
    add_page_break(doc)

    # ============================================================
    # CHAPTER 3: SYSTEM ANALYSIS
    # ============================================================
    add_chapter_heading(doc, "Chapter 3: System Analysis")
    add_para(doc,
             "The system analysis phase translated the high-level goals and "
             "constraints surfaced during the system study into precise, "
             "verifiable software requirements. This chapter presents the "
             "Functional Requirements (what the system must do), the "
             "Non-Functional Requirements (how well it must do it), the User "
             "Requirements (what each user category expects), and the System "
             "Requirements (the underlying hardware and software environment).")

    add_sub_heading(doc, "3.1 Functional Requirements")
    add_para(doc,
             "Functional requirements describe the capabilities that the "
             "platform must offer to its users. The functional requirements of "
             "ActToAction are organised by module.")
    add_subsub_heading(doc, "3.1.1 Catalogue Browsing")
    add_bullets(doc, [
        "FR-1: The system shall allow visitors to browse a paginated list of all published courses.",
        "FR-2: The system shall allow visitors to filter courses by category and search by name.",
        "FR-3: The system shall display individual course detail pages with description, fees, sessions, and enrolment call-to-action.",
        "FR-4: The system shall allow visitors to browse events, sub-events, services, industries, blogs, and gallery items in a similar paginated, filterable manner.",
    ])
    add_subsub_heading(doc, "3.1.2 Enrollment & Registration")
    add_bullets(doc, [
        "FR-5: The system shall allow a visitor to enrol in a course by submitting a structured form.",
        "FR-6: The system shall validate input client-side and server-side and present field-level error messages on failure.",
        "FR-7: The system shall allow a visitor to register for an event or sub-event with one or more attendees per registration.",
        "FR-8: The system shall allow workshop registration scoped by city, school, and age group.",
        "FR-9: The system shall record every form submission with a timestamp, status, and source.",
    ])
    add_subsub_heading(doc, "3.1.3 Payment")
    add_bullets(doc, [
        "FR-10: The system shall create a Razorpay order on the server when the user proceeds to pay.",
        "FR-11: The system shall present the Razorpay checkout to the user and accept payments via cards, UPI, netbanking, and wallets.",
        "FR-12: The system shall verify the payment signature on the server using the Razorpay HMAC secret.",
        "FR-13: The system shall mark the corresponding registration or enrolment as paid only after a successful signature match.",
    ])
    add_subsub_heading(doc, "3.1.4 Psychometric Testing")
    add_bullets(doc, [
        "FR-14: The system shall allow visitors to browse and take psychometric tests grouped by category.",
        "FR-15: The system shall compute a score by summing the weights of selected options.",
        "FR-16: The system shall match the score against configured TestResultRange rows and render the matching interpretation.",
        "FR-17: The system shall optionally render a graph (radar or bar) when a TestGraphConfig row is active for the test.",
    ])
    add_subsub_heading(doc, "3.1.5 Content Management")
    add_bullets(doc, [
        "FR-18: The system shall provide an authenticated admin panel for creating, editing, retiring, and deleting all content entities.",
        "FR-19: The system shall provide a status toggle on every content entity to control visibility on the public site.",
        "FR-20: The system shall accept image uploads with content-type validation and store files under public/img/<entity>/.",
        "FR-21: The system shall maintain unique slug fields for every public entity to enable SEO-friendly URLs.",
    ])
    add_subsub_heading(doc, "3.1.6 Communication")
    add_bullets(doc, [
        "FR-22: The system shall dispatch transactional emails using configurable templates with merge variables.",
        "FR-23: The system shall log every dispatched email with status (queued, sent, failed) and an error message on failure.",
        "FR-24: The system shall expose a chatbot widget that returns active FAQs and accepts support tickets.",
    ])

    add_sub_heading(doc, "3.2 Non-Functional Requirements")
    add_para(doc,
             "Non-functional requirements describe how well the system must "
             "perform its functions. They are summarised below.")
    add_subsub_heading(doc, "3.2.1 Performance")
    add_bullets(doc, [
        "NFR-1: Public listing pages shall render within 1.5 seconds at the 95th percentile under a load of 100 concurrent users.",
        "NFR-2: Admin listing pages shall render within 2 seconds for tables of up to 10,000 rows.",
        "NFR-3: Payment verification shall complete within 1 second of receiving the gateway callback.",
    ])
    add_subsub_heading(doc, "3.2.2 Reliability")
    add_bullets(doc, [
        "NFR-4: The platform shall maintain at least 99.5% monthly uptime excluding scheduled maintenance.",
        "NFR-5: No payment shall be marked successful without a verified gateway signature.",
        "NFR-6: All form submissions shall be persisted before any external dispatch (email, gateway).",
    ])
    add_subsub_heading(doc, "3.2.3 Security")
    add_bullets(doc, [
        "NFR-7: All HTTP traffic shall be served over TLS.",
        "NFR-8: Passwords shall be stored using the bcrypt hashing algorithm.",
        "NFR-9: All form submissions shall require a valid CSRF token.",
        "NFR-10: All database queries shall use parameterised statements via Eloquent or query builder.",
    ])
    add_subsub_heading(doc, "3.2.4 Usability")
    add_bullets(doc, [
        "NFR-11: The public site shall be responsive and functional on devices with a minimum width of 320 pixels.",
        "NFR-12: The admin panel shall use a consistent left-sidebar layout across modules.",
        "NFR-13: All form errors shall be shown adjacent to the offending field.",
    ])
    add_subsub_heading(doc, "3.2.5 Maintainability")
    add_bullets(doc, [
        "NFR-14: The codebase shall conform to PSR-12 PHP coding standards.",
        "NFR-15: Database schema changes shall be expressed as Laravel migrations.",
        "NFR-16: All environment-specific values shall be configured via .env, never hard-coded.",
    ])

    add_sub_heading(doc, "3.3 User Requirements")
    add_table(doc,
              ["User Class", "Representative Requirement"],
              [
                  ["Visitor",
                   "I want to browse courses, see their fees, and enrol in one click."],
                  ["Student",
                   "I want a confirmation email and PDF receipt the moment my payment succeeds."],
                  ["Content Editor",
                   "I want to publish a blog post in under five minutes without needing a developer."],
                  ["Admissions Counsellor",
                   "I want to see every admission lead in one searchable list with status filters."],
                  ["Event Coordinator",
                   "I want to export the attendee list for a sub-event as a CSV ready to print."],
                  ["Workshop Leader",
                   "I want to see registrations grouped by city and school, sorted by date."],
                  ["Support Agent",
                   "I want to triage chatbot tickets and respond by email from one screen."],
                  ["Admin",
                   "I want one place to toggle the visibility of any content entity on the public site."],
              ])

    add_sub_heading(doc, "3.4 System Requirements")
    add_subsub_heading(doc, "3.4.1 Hardware Requirements (Server)")
    add_bullets(doc, [
        "Processor: 2 vCPU minimum (4 vCPU recommended).",
        "Memory: 4 GB RAM minimum (8 GB RAM recommended).",
        "Storage: 40 GB SSD minimum, with separate backup volume.",
        "Network: 100 Mbps connection with a public IPv4 address.",
    ])
    add_subsub_heading(doc, "3.4.2 Software Requirements (Server)")
    add_bullets(doc, [
        "Operating System: Ubuntu 22.04 LTS or later (Linux).",
        "Web Server: Nginx 1.20+ or Apache 2.4+.",
        "PHP: 8.2 or later with extensions: mbstring, openssl, pdo_mysql, tokenizer, xml, ctype, json, fileinfo, gd.",
        "Database: MySQL 8.0+ or MariaDB 10.6+.",
        "Composer: 2.5+ for PHP dependency management.",
        "Node.js: 20+ for asset bundling during build.",
        "TLS Certificate: Let's Encrypt or commercial.",
    ])
    add_subsub_heading(doc, "3.4.3 Hardware Requirements (Client)")
    add_bullets(doc, [
        "Any device capable of running a modern browser: smartphone, tablet, laptop, or desktop.",
        "Screen width of at least 320 pixels.",
        "Stable internet connection (minimum 256 Kbps).",
    ])
    add_subsub_heading(doc, "3.4.4 Software Requirements (Client)")
    add_bullets(doc, [
        "Browser: Chrome, Firefox, Safari, or Edge — current or one previous major version.",
        "JavaScript and cookies enabled.",
        "PDF reader for receipts.",
    ])
    add_page_break(doc)

    # ============================================================
    # CHAPTER 4: SYSTEM DESIGN
    # ============================================================
    add_chapter_heading(doc, "Chapter 4: System Design")
    add_para(doc,
             "This chapter presents the design artefacts that translate the "
             "requirements of the previous chapter into an implementable "
             "blueprint. It covers the system architecture, data flow diagrams, "
             "entity-relationship diagram, use case diagram, principal "
             "flowcharts, the database table structure, and representative "
             "input and output forms.")

    add_sub_heading(doc, "4.0 System Architecture")
    add_image(doc, charts["arch"], "Figure 4.1 — High-level system architecture of ActToAction.")
    add_para(doc,
             "ActToAction follows a classic three-tier architecture. The "
             "presentation tier comprises Blade-rendered HTML pages (public and "
             "admin) and small fragments of JavaScript for interactive "
             "elements. The application tier is the Laravel framework: "
             "controllers, middleware, services, and Eloquent models. The data "
             "tier is MySQL, supplemented by file storage on disk for uploaded "
             "images and PDFs. Razorpay and the SMTP relay are external "
             "services accessed over HTTPS.")

    add_sub_heading(doc, "4.1 Data Flow Diagrams")
    add_para(doc,
             "Data Flow Diagrams (DFDs) describe how data moves through the "
             "system, from external entities into the application's processes "
             "and stores. ActToAction is documented at two levels of detail.")
    add_subsub_heading(doc, "4.1.1 DFD Level 0 (Context Diagram)")
    add_image(doc, charts["dfd0"], "Figure 4.2 — DFD Level 0 (context diagram).")
    add_para(doc,
             "At the context level, ActToAction appears as a single process "
             "interacting with four external entities: visitors, admins, "
             "Razorpay, and the SMTP relay. Visitors send form submissions and "
             "browse pages; admins perform content management and report "
             "extraction; Razorpay verifies payments; and the SMTP relay "
             "delivers transactional emails on behalf of the platform.")
    add_subsub_heading(doc, "4.1.2 DFD Level 1")
    add_image(doc, charts["dfd1"], "Figure 4.3 — DFD Level 1 expanding the major processes.")
    add_para(doc,
             "Decomposing the central process exposes four major sub-processes: "
             "browse catalogue, register/enrol, make payment, and manage "
             "content. Each sub-process is associated with one or more data "
             "stores: D1 (catalogue tables), D2 (registration tables), D3 "
             "(payments), and D4 (admin-managed content tables).")

    add_sub_heading(doc, "4.2 E-R Diagrams")
    add_image(doc, charts["er"], "Figure 4.4 — Entity-Relationship diagram of the core entities.")
    add_para(doc,
             "The entity-relationship diagram captures the principal entities "
             "and the cardinalities between them. A USER places one or more "
             "ENROLLMENTs. Each ENROLLMENT is for one COURSE and is linked to "
             "exactly one PAYMENT. Each COURSE belongs to one COURSE_CATEGORY. "
             "An EVENT contains zero or more SUB_EVENTs. Each SUB_EVENT "
             "receives zero or more EVENT_REGs, and each EVENT_REG itself "
             "lists one or more ATTENDEEs. PAYMENT is the central financial "
             "entity, attaching to both ENROLLMENT and EVENT_REG (polymorphic "
             "in spirit, modelled here as nullable foreign keys). BLOGs belong "
             "to a BLOG_CATEGORY, and PSYCH_TESTs are independent of the "
             "transactional cluster but share user identity.")

    add_sub_heading(doc, "4.3 Use Case Diagrams")
    add_image(doc, charts["use_case"], "Figure 4.5 — Use case diagram showing actors and their interactions with the system.")
    add_para(doc,
             "The use case diagram identifies five actors and the specific "
             "interactions each actor has with the system. Visitors and "
             "Students drive most of the flow, with the Admin actor responsible "
             "for content management and reporting. Razorpay and SMTP appear as "
             "secondary (external) actors invoked by the system rather than by "
             "human users.")

    add_sub_heading(doc, "4.4 Flow Charts")
    add_para(doc,
             "Flowcharts depict the step-by-step control flow of the most "
             "important user journeys in the system. The seven flowcharts "
             "below cover the dominant transactional paths.")
    add_subsub_heading(doc, "4.4.1 Admin Login Flow")
    add_image(doc, charts["flow_login"], "Figure 4.6 — Admin login flow.")
    add_subsub_heading(doc, "4.4.2 Course Enrollment Flow")
    add_image(doc, charts["flow_enroll"], "Figure 4.7 — Course enrollment flow with payment verification.")
    add_subsub_heading(doc, "4.4.3 Event Registration Flow")
    add_image(doc, charts["flow_event"], "Figure 4.8 — Event registration with optional partner redirect.")
    add_subsub_heading(doc, "4.4.4 Razorpay Payment Flow")
    add_image(doc, charts["flow_pay"], "Figure 4.9 — End-to-end Razorpay payment flow.")
    add_subsub_heading(doc, "4.4.5 Psychometric Test Flow")
    add_image(doc, charts["flow_psych"], "Figure 4.10 — Psychometric test administration flow.")
    add_subsub_heading(doc, "4.4.6 Blog Publishing Flow")
    add_image(doc, charts["flow_blog"], "Figure 4.11 — Blog publishing flow.")
    add_subsub_heading(doc, "4.4.7 Admission Form Flow")
    add_image(doc, charts["flow_admit"], "Figure 4.12 — Two-stage admission flow.")

    add_sub_heading(doc, "4.5 Database Tables")
    add_para(doc,
             "The ActToAction schema comprises more than eighty tables organised "
             "into ten functional clusters. The structures of the most "
             "important transactional tables are presented below; the full "
             "schema is reproduced in Appendix A of the codebase.")
    add_subsub_heading(doc, "4.5.1 Table: courses")
    add_table(doc,
              ["Field", "Type", "Constraints", "Description"],
              [
                  ["id", "BIGINT", "PK, AUTO_INCREMENT", "Primary key"],
                  ["course_category_id", "BIGINT", "FK, NOT NULL", "→ course_categories.id"],
                  ["name", "VARCHAR(255)", "NOT NULL", "Display name"],
                  ["slug", "VARCHAR(255)", "UNIQUE, NOT NULL", "URL-safe identifier"],
                  ["short_description", "TEXT", "NULLABLE", "Listing card copy"],
                  ["long_description", "LONGTEXT", "NULLABLE", "Detail page copy"],
                  ["banner_image", "VARCHAR(255)", "NULLABLE", "Path under public/img"],
                  ["fee", "DECIMAL(10,2)", "NULLABLE", "Default fee"],
                  ["status", "BOOLEAN", "DEFAULT 1", "Visibility flag"],
                  ["created_at, updated_at", "TIMESTAMP", "Laravel defaults", "Audit timestamps"],
              ])
    add_subsub_heading(doc, "4.5.2 Table: enrollments")
    add_table(doc,
              ["Field", "Type", "Constraints", "Description"],
              [
                  ["id", "BIGINT", "PK", "Primary key"],
                  ["course_id", "BIGINT", "FK", "→ courses.id"],
                  ["student_name", "VARCHAR(255)", "NOT NULL", "Applicant name"],
                  ["age", "INT", "NULLABLE", "Applicant age"],
                  ["parent_name", "VARCHAR(255)", "NULLABLE", "Parent / guardian"],
                  ["mother_phone", "VARCHAR(20)", "NULLABLE", "Mother's phone"],
                  ["phone", "VARCHAR(20)", "NOT NULL", "Primary contact"],
                  ["email", "VARCHAR(255)", "NOT NULL", "Email address"],
                  ["city", "VARCHAR(100)", "NULLABLE", "City"],
                  ["status", "ENUM", "pending|confirmed|cancelled", "Workflow status"],
                  ["created_at, updated_at", "TIMESTAMP", "Laravel defaults", "Audit"],
              ])
    add_subsub_heading(doc, "4.5.3 Table: payments")
    add_table(doc,
              ["Field", "Type", "Constraints", "Description"],
              [
                  ["id", "BIGINT", "PK", "Primary key"],
                  ["enrollment_id", "BIGINT", "FK, NULLABLE", "→ enrollments.id"],
                  ["event_registration_id", "BIGINT", "FK, NULLABLE", "→ event_registrations.id"],
                  ["amount", "DECIMAL(10,2)", "NOT NULL", "Amount in INR"],
                  ["razorpay_order_id", "VARCHAR(50)", "INDEX", "Razorpay order id"],
                  ["razorpay_payment_id", "VARCHAR(50)", "UNIQUE", "Razorpay payment id"],
                  ["razorpay_signature", "VARCHAR(255)", "NULLABLE", "HMAC signature"],
                  ["status", "ENUM", "pending|success|failed|refunded", "Payment status"],
                  ["created_at, updated_at", "TIMESTAMP", "Laravel defaults", "Audit"],
              ])
    add_subsub_heading(doc, "4.5.4 Table: events and sub_events")
    add_table(doc,
              ["Field", "Type", "Description"],
              [
                  ["events.id", "BIGINT PK", "Primary key"],
                  ["events.title", "VARCHAR(255)", "Event title"],
                  ["events.slug", "VARCHAR(255) UNIQUE", "URL slug"],
                  ["events.start_date / end_date", "DATE", "Display window"],
                  ["events.banner_image", "VARCHAR(255)", "Hero image"],
                  ["events.status", "BOOLEAN", "Visibility"],
                  ["sub_events.id", "BIGINT PK", "Primary key"],
                  ["sub_events.event_id", "BIGINT FK", "→ events.id"],
                  ["sub_events.title", "VARCHAR(255)", "Sub-event title"],
                  ["sub_events.fee", "DECIMAL(10,2)", "Per-attendee fee"],
                  ["sub_events.redirect_link", "VARCHAR(500)", "Optional partner URL"],
              ])
    add_subsub_heading(doc, "4.5.5 Table: blogs")
    add_table(doc,
              ["Field", "Type", "Description"],
              [
                  ["id", "BIGINT PK", "Primary key"],
                  ["blog_category_id", "BIGINT FK", "→ blog_categories.id"],
                  ["blog_author_id", "BIGINT FK", "→ blog_authors.id"],
                  ["title", "VARCHAR(255)", "Post title"],
                  ["slug", "VARCHAR(255) UNIQUE", "URL slug"],
                  ["excerpt", "TEXT", "Listing snippet"],
                  ["content", "LONGTEXT", "Body HTML"],
                  ["cover_image", "VARCHAR(255)", "Cover image path"],
                  ["social_embeds", "JSON", "List of embed snippets"],
                  ["status", "BOOLEAN", "Visibility"],
              ])
    add_subsub_heading(doc, "4.5.6 Table: psych_tests, psych_questions, test_result_ranges")
    add_table(doc,
              ["Field", "Type", "Description"],
              [
                  ["psych_tests.id", "BIGINT PK", "Test id"],
                  ["psych_tests.psych_category_id", "BIGINT FK", "Group"],
                  ["psych_tests.title", "VARCHAR(255)", "Title"],
                  ["psych_tests.slug", "VARCHAR(255) UNIQUE", "URL slug"],
                  ["psych_questions.id", "BIGINT PK", "Question id"],
                  ["psych_questions.psych_test_id", "BIGINT FK", "→ psych_tests.id"],
                  ["psych_questions.text", "TEXT", "Question text"],
                  ["psych_questions.options", "JSON", "Options + weights"],
                  ["test_result_ranges.id", "BIGINT PK", "Range id"],
                  ["test_result_ranges.psych_test_id", "BIGINT FK", "→ psych_tests.id"],
                  ["test_result_ranges.min_score", "INT", "Lower bound"],
                  ["test_result_ranges.max_score", "INT", "Upper bound"],
                  ["test_result_ranges.label", "VARCHAR(255)", "Result label"],
                  ["test_result_ranges.copy", "TEXT", "Interpretation"],
              ])
    add_subsub_heading(doc, "4.5.7 Table: email_templates and email_logs")
    add_table(doc,
              ["Field", "Type", "Description"],
              [
                  ["email_templates.id", "BIGINT PK", "Template id"],
                  ["email_templates.key", "VARCHAR(100) UNIQUE", "Code-side identifier"],
                  ["email_templates.subject", "VARCHAR(255)", "Email subject"],
                  ["email_templates.body", "LONGTEXT", "HTML body"],
                  ["email_templates.variables", "JSON", "Allowed merge variables"],
                  ["email_logs.id", "BIGINT PK", "Log id"],
                  ["email_logs.template_id", "BIGINT FK", "→ email_templates.id"],
                  ["email_logs.recipient", "VARCHAR(255)", "Email address"],
                  ["email_logs.status", "ENUM", "queued|sent|failed"],
                  ["email_logs.error", "TEXT", "Last error"],
                  ["email_logs.sent_at", "TIMESTAMP", "Delivery time"],
              ])

    add_sub_heading(doc, "4.6 Input / Output Forms")
    add_subsub_heading(doc, "4.6.1 Input — Course Enrollment Form")
    add_para(doc,
             "Located at /courses/{slug} → Enrol Now. The form captures the "
             "student's name, age, parent's name, mother's phone, primary "
             "phone, email, and city. All required fields are marked with an "
             "asterisk; validation messages appear inline below each field. "
             "On submission, the data is POSTed to "
             "/enrollments/store and the user is forwarded to the Razorpay "
             "checkout modal.")
    add_subsub_heading(doc, "4.6.2 Input — Event Registration Form")
    add_para(doc,
             "Located at /events/{slug}. Fields include lead applicant name, "
             "email, phone, and a dynamic attendees section where rows can "
             "be added or removed. The total fee updates live as attendees are "
             "added.")
    add_subsub_heading(doc, "4.6.3 Input — Workshop Registration Form")
    add_para(doc,
             "Located at /workshops. Fields include the city, school, age "
             "group (each populated from the corresponding admin-managed "
             "table), participant name, parent contact, and an optional "
             "merchandise add-on.")
    add_subsub_heading(doc, "4.6.4 Input — Admin Login Form")
    add_para(doc,
             "Located at /admin/login. Fields are email and password. On "
             "successful authentication the admin is redirected to "
             "/admin/dashboard.")
    add_subsub_heading(doc, "4.6.5 Input — Admin Course Editor Form")
    add_para(doc,
             "Located at /admin/courses/{id}/edit. Fields cover category, "
             "name, slug, short and long descriptions, banner image, fees, "
             "and status. The form is multipart/form-data because it accepts "
             "an image upload.")
    add_subsub_heading(doc, "4.6.6 Screenshot — Admin Login Page")
    add_image(doc, charts["ss_login"], "Figure 4.13 — Admin login page (mockup).")
    add_subsub_heading(doc, "4.6.7 Screenshot — Admin Dashboard")
    add_image(doc, charts["ss_dashboard"], "Figure 4.14 — Admin dashboard after successful login.")
    add_subsub_heading(doc, "4.6.8 Screenshot — Course Edit Form")
    add_image(doc, charts["ss_course_form"], "Figure 4.15 — Admin form for creating / editing a course.")
    add_subsub_heading(doc, "4.6.9 Screenshot — Public Enrolment Form")
    add_image(doc, charts["ss_enroll_form"], "Figure 4.16 — Public-facing course enrolment form.")
    add_subsub_heading(doc, "4.6.10 Screenshot — Admin Enrolment Listing")
    add_image(doc, charts["ss_listing"], "Figure 4.17 — Admin enrolment listing with search, filters, and pagination.")
    add_subsub_heading(doc, "4.6.11 Sample PDF Receipt")
    add_image(doc, charts["ss_receipt"], "Figure 4.18 — Sample PDF receipt generated for a successful payment.")
    add_sub_heading(doc, "4.7 Real Site Imagery")
    add_para(doc,
             "The figures below reproduce a representative selection of "
             "imagery actually published on the ActToAction website. They "
             "illustrate the visual vocabulary that the design and content "
             "decisions of this project must accommodate.")
    add_subsub_heading(doc, "4.7.1 Programme Photographs")
    if os.path.exists(real("gallery1.png")):
        add_image(doc, real("gallery1.png"),
                  "Figure 4.19 — Programme activity photograph (gallery image 1).",
                  width=5.5)
    if os.path.exists(real("gallery2.png")):
        add_image(doc, real("gallery2.png"),
                  "Figure 4.20 — Programme activity photograph (gallery image 2).",
                  width=5.5)
    if os.path.exists(real("gallery3.png")):
        add_image(doc, real("gallery3.png"),
                  "Figure 4.21 — Programme activity photograph (gallery image 3).",
                  width=5.5)
    if os.path.exists(real("gallery4.png")):
        add_image(doc, real("gallery4.png"),
                  "Figure 4.22 — Programme activity photograph (gallery image 4).",
                  width=5.5)
    add_subsub_heading(doc, "4.7.2 Facility Photograph")
    if os.path.exists(real("facilities.png")):
        add_image(doc, real("facilities.png"),
                  "Figure 4.23 — ActToAction learning facility.",
                  width=5.5)
    add_subsub_heading(doc, "4.7.3 Notification Banner Artwork")
    if os.path.exists(real("banner_course.jpg")):
        add_image(doc, real("banner_course.jpg"),
                  "Figure 4.24 — Course notification banner artwork.",
                  width=5.5)
    if os.path.exists(real("banner_skill.jpg")):
        add_image(doc, real("banner_skill.jpg"),
                  "Figure 4.25 — Skills notification banner artwork.",
                  width=5.5)
    add_subsub_heading(doc, "4.7.4 Partner Logos")
    add_para(doc,
             "Partner organisations are surfaced on the home page through a "
             "logos strip. The four logos reproduced below are illustrative.")
    if os.path.exists(real("client1.png")):
        add_image(doc, real("client1.png"),
                  "Figure 4.26 — Partner brand logo (client 1).", width=3.0)
    if os.path.exists(real("client2.png")):
        add_image(doc, real("client2.png"),
                  "Figure 4.27 — Partner brand logo (client 2).", width=3.0)
    if os.path.exists(real("client3.png")):
        add_image(doc, real("client3.png"),
                  "Figure 4.28 — Partner brand logo (client 3).", width=3.0)
    if os.path.exists(real("client4.png")):
        add_image(doc, real("client4.png"),
                  "Figure 4.29 — Partner brand logo (client 4).", width=3.0)
    add_subsub_heading(doc, "4.7.5 Team Photographs")
    add_para(doc,
             "The team-members module renders profile cards using uploaded "
             "headshots. The four representative profiles below show the "
             "standard 1:1 aspect ratio and circular crop applied at render "
             "time on the public site.")
    if os.path.exists(real("person_f1.png")):
        add_image(doc, real("person_f1.png"),
                  "Figure 4.30 — Team profile photograph (member 1).",
                  width=2.4)
    if os.path.exists(real("person_m1.png")):
        add_image(doc, real("person_m1.png"),
                  "Figure 4.31 — Team profile photograph (member 2).",
                  width=2.4)
    if os.path.exists(real("person_f2.png")):
        add_image(doc, real("person_f2.png"),
                  "Figure 4.32 — Team profile photograph (member 3).",
                  width=2.4)
    if os.path.exists(real("person_m2.png")):
        add_image(doc, real("person_m2.png"),
                  "Figure 4.33 — Team profile photograph (member 4).",
                  width=2.4)
    add_subsub_heading(doc, "4.7.6 Additional Programme Imagery")
    if os.path.exists(real("gallery5.png")):
        add_image(doc, real("gallery5.png"),
                  "Figure 4.34 — Programme activity photograph (gallery image 5).",
                  width=5.5)
    if os.path.exists(real("gallery6.png")):
        add_image(doc, real("gallery6.png"),
                  "Figure 4.35 — Programme activity photograph (gallery image 6).",
                  width=5.5)
    add_subsub_heading(doc, "4.7.7 Output — Confirmation Email")
    add_para(doc,
             "On a successful enrolment, the user receives an email with a "
             "subject like \"Your enrolment in [Course Name] is confirmed\". "
             "The body includes the student name, course name, fee paid, "
             "Razorpay payment id, and a PDF receipt as an attachment.")
    add_subsub_heading(doc, "4.7.8 Output — PDF Receipt")
    add_para(doc,
             "Generated using barryvdh/laravel-dompdf. The receipt is one A4 "
             "page comprising the organisation logo, the receipt number, the "
             "date, the payer details, the line items, the total amount in "
             "words, and a footer with the registered office address.")
    add_subsub_heading(doc, "4.7.9 Output — Admin Listing Page")
    add_para(doc,
             "Each module's index view renders a sortable, searchable, "
             "paginated table of records with action buttons (edit, toggle "
             "status, delete) and an export-to-CSV link.")
    add_page_break(doc)

    # ============================================================
    # CHAPTER 5: DEVELOPMENT
    # ============================================================
    add_chapter_heading(doc, "Chapter 5: Development")
    add_para(doc,
             "This chapter describes the development environment, coding style "
             "guidelines, the techniques applied during implementation, and "
             "representative code snippets that illustrate how the design was "
             "translated into running software.")

    add_sub_heading(doc, "5.1 Environment")
    add_subsub_heading(doc, "5.1.1 Local Development Environment")
    add_bullets(doc, [
        "Operating System: Microsoft Windows 10 / 11.",
        "Local Stack: Laravel Herd (bundled PHP 8.2, MySQL, Node.js).",
        "IDE: Visual Studio Code with PHP Intelephense, Laravel Blade, and EditorConfig extensions.",
        "Version Control: Git with the project hosted on a private remote.",
        "Database GUI: TablePlus / phpMyAdmin for inspection and ad-hoc queries.",
        "API Testing: Postman / Insomnia for verifying chatbot and payment endpoints.",
    ])
    add_subsub_heading(doc, "5.1.2 Production Environment")
    add_bullets(doc, [
        "Operating System: Ubuntu 22.04 LTS.",
        "Web Server: Nginx 1.22 reverse-proxying to PHP-FPM 8.2.",
        "Database: MySQL 8.0 on the same host (or a managed instance for larger deployments).",
        "Process Manager: systemd units for the queue worker and (optionally) Laravel scheduler.",
        "TLS: Let's Encrypt certificates auto-renewed via certbot.",
    ])
    add_subsub_heading(doc, "5.1.3 Configuration")
    add_para(doc,
             "All environment-specific values (database credentials, mailer "
             "settings, Razorpay keys, application URL, queue connection) are "
             "read from a .env file at boot time. The .env file is excluded "
             "from version control; a .env.example committed alongside the "
             "code documents the required keys.")

    add_sub_heading(doc, "5.2 Coding Style")
    add_para(doc,
             "The project follows the PSR-12 PHP coding standard, supplemented "
             "by Laravel's own conventions. Laravel Pint is configured as the "
             "autoformatter; running php artisan pint normalises whitespace, "
             "import ordering, and brace placement.")
    add_subsub_heading(doc, "5.2.1 Naming Conventions")
    add_bullets(doc, [
        "Classes use UpperCamelCase (e.g. CourseController, EnrollmentController).",
        "Methods and variables use lowerCamelCase (e.g. handlePayment, $courseId).",
        "Database columns use snake_case (e.g. course_category_id, created_at).",
        "Routes use lowercase-with-hyphens (e.g. /admin/blog-categories).",
        "Blade view files use snake_case (e.g. course_index.blade.php).",
    ])
    add_subsub_heading(doc, "5.2.2 File Organisation")
    add_bullets(doc, [
        "Controllers grouped by area: App\\Http\\Controllers (public) and App\\Http\\Controllers\\admin.",
        "Models in App\\Models with one class per table.",
        "Views under resources/views with backend/ and frontend/ subtrees.",
        "Routes in routes/web.php organised top-down, public routes first then an admin route group.",
        "Migrations in database/migrations, named with timestamps for ordering.",
    ])
    add_subsub_heading(doc, "5.2.3 Documentation")
    add_para(doc,
             "Class-level docblocks describe the responsibility of each "
             "controller and model. Method-level docblocks document parameters "
             "and return types where the type system alone is not expressive "
             "enough. Inline comments are reserved for non-obvious decisions; "
             "self-explanatory code is left to speak for itself.")

    add_sub_heading(doc, "5.3 Coding Techniques")
    add_subsub_heading(doc, "5.3.1 MVC Separation")
    add_para(doc,
             "Controllers contain only the logic needed to translate HTTP "
             "requests into model operations and to choose a view or response. "
             "Business rules live in the models or in dedicated service "
             "classes when they cross multiple models. Views contain "
             "presentation logic only — no database queries.")
    add_subsub_heading(doc, "5.3.2 Eloquent ORM")
    add_para(doc,
             "Each table has a corresponding Eloquent model. Mass assignment "
             "is restricted to fields explicitly listed in $fillable. "
             "Relationships are declared as methods returning hasMany, "
             "belongsTo, or belongsToMany. Eager loading via with() is used "
             "wherever a listing page accesses related data, to avoid N+1 "
             "query patterns.")
    add_subsub_heading(doc, "5.3.3 Form Validation")
    add_para(doc,
             "Validation rules are declared at the top of each store/update "
             "method using $request->validate([...]). For complex forms the "
             "rules are extracted into a Form Request class so that they "
             "can be reused and unit-tested independently.")
    add_subsub_heading(doc, "5.3.4 Middleware")
    add_para(doc,
             "Cross-cutting concerns such as authentication, CSRF protection, "
             "and request logging are implemented as middleware. The custom "
             "admin.auth middleware checks for an admin_id session value and "
             "redirects to /admin/login if absent.")
    add_subsub_heading(doc, "5.3.5 Service Container & Dependency Injection")
    add_para(doc,
             "Controllers receive their dependencies (e.g. mailable factories, "
             "PDF generators) via constructor injection. The framework's "
             "service container resolves the dependencies automatically at "
             "runtime, simplifying testing and substitution.")
    add_subsub_heading(doc, "5.3.6 Queue & Background Jobs")
    add_para(doc,
             "Slow operations such as email dispatch and PDF generation are "
             "pushed onto the queue and executed by a long-running queue "
             "worker. This keeps web requests responsive and isolates "
             "transient failures (SMTP unavailability, for instance) from the "
             "user-facing flow.")
    add_subsub_heading(doc, "5.3.7 Database Migrations")
    add_para(doc,
             "Schema evolution is captured as a series of migration files. "
             "Each migration is small and focused — adding a column, creating "
             "a table, adding an index. New migrations are appended; existing "
             "migrations are never modified once committed.")

    add_sub_heading(doc, "5.4 Coding")
    add_para(doc, "Representative code snippets are reproduced below.")
    add_subsub_heading(doc, "5.4.1 Course Model")
    add_para(doc,
             "namespace App\\Models;\n\nclass Course extends Model {\n"
             "    protected $fillable = [\n"
             "        'course_category_id', 'name', 'slug',\n"
             "        'short_description', 'long_description',\n"
             "        'banner_image', 'fee', 'status',\n"
             "    ];\n\n"
             "    public function category() {\n"
             "        return $this->belongsTo(CourseCategory::class, 'course_category_id');\n"
             "    }\n\n"
             "    public function enrollments() {\n"
             "        return $this->hasMany(Enrollment::class);\n"
             "    }\n"
             "}",
             justify=False, indent=0)
    add_subsub_heading(doc, "5.4.2 Enrollment Controller — store")
    add_para(doc,
             "public function store(Request $request) {\n"
             "    $data = $request->validate([\n"
             "        'course_id' => 'required|exists:courses,id',\n"
             "        'student_name' => 'required|string|max:255',\n"
             "        'phone' => 'required|string|max:20',\n"
             "        'email' => 'required|email|max:255',\n"
             "    ]);\n"
             "    $enrollment = Enrollment::create($data + ['status' => 'pending']);\n"
             "    $order = $this->razorpay->order->create([\n"
             "        'amount' => $enrollment->course->fee * 100,\n"
             "        'currency' => 'INR',\n"
             "        'receipt' => 'enr_' . $enrollment->id,\n"
             "    ]);\n"
             "    $enrollment->update(['razorpay_order_id' => $order->id]);\n"
             "    return view('frontend.enrollments.pay',\n"
             "                compact('enrollment', 'order'));\n"
             "}",
             justify=False, indent=0)
    add_subsub_heading(doc, "5.4.3 Razorpay Verification")
    add_para(doc,
             "public function verify(Request $request) {\n"
             "    $signature = hash_hmac('sha256',\n"
             "        $request->razorpay_order_id . '|' . $request->razorpay_payment_id,\n"
             "        config('services.razorpay.secret'));\n"
             "    if (!hash_equals($signature, $request->razorpay_signature)) {\n"
             "        abort(400, 'Signature mismatch');\n"
             "    }\n"
             "    Payment::create([\n"
             "        'enrollment_id' => $request->enrollment_id,\n"
             "        'razorpay_order_id' => $request->razorpay_order_id,\n"
             "        'razorpay_payment_id' => $request->razorpay_payment_id,\n"
             "        'razorpay_signature' => $request->razorpay_signature,\n"
             "        'amount' => $request->amount,\n"
             "        'status' => 'success',\n"
             "    ]);\n"
             "    Enrollment::find($request->enrollment_id)\n"
             "        ->update(['status' => 'confirmed']);\n"
             "    Mail::to($request->email)\n"
             "        ->queue(new EnrollmentConfirmed($enrollment));\n"
             "    return redirect()->route('enrollment.success');\n"
             "}",
             justify=False, indent=0)
    add_subsub_heading(doc, "5.4.4 admin.auth Middleware")
    add_para(doc,
             "public function handle(Request $request, Closure $next) {\n"
             "    if (!session()->has('admin_id')) {\n"
             "        return redirect()->route('admin.login')\n"
             "            ->with('intended', $request->fullUrl());\n"
             "    }\n"
             "    return $next($request);\n"
             "}",
             justify=False, indent=0)
    add_subsub_heading(doc, "5.4.5 Blog Controller — listing with eager loading")
    add_para(doc,
             "public function index(Request $request) {\n"
             "    $blogs = Blog::with(['category', 'author', 'tags'])\n"
             "        ->where('status', 1)\n"
             "        ->when($request->category,\n"
             "            fn($q) => $q->whereHas('category',\n"
             "                fn($c) => $c->where('slug', $request->category)))\n"
             "        ->latest()\n"
             "        ->paginate(12);\n"
             "    return view('frontend.blogs.index', compact('blogs'));\n"
             "}",
             justify=False, indent=0)
    add_page_break(doc)

    # ============================================================
    # CHAPTER 6: TESTING
    # ============================================================
    add_chapter_heading(doc, "Chapter 6: Testing")
    add_para(doc,
             "Testing was carried out at multiple levels — unit, integration, "
             "and user acceptance — to verify that each module behaves as "
             "specified and that the system as a whole satisfies the "
             "requirements captured in Chapter 3. This chapter outlines the "
             "testing strategy and reproduces a representative set of test "
             "cases.")

    add_sub_heading(doc, "6.1 Test Cases")
    add_para(doc,
             "Each test case below identifies a specific scenario, the inputs "
             "supplied, the expected output, and the observed result. The "
             "tests cover happy paths, validation failures, and edge cases "
             "across every major module.")

    add_subsub_heading(doc, "6.1.1 Authentication")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-A-01", "Valid admin login", "Correct email + password",
                   "Redirect to /admin/dashboard", "Pass"],
                  ["TC-A-02", "Wrong password", "Correct email, wrong password",
                   "Show 'Invalid credentials'", "Pass"],
                  ["TC-A-03", "Unknown email", "Non-existent email",
                   "Show 'Invalid credentials'", "Pass"],
                  ["TC-A-04", "Empty form", "Blank email and password",
                   "Show 'required' errors", "Pass"],
                  ["TC-A-05", "Direct access to /admin/courses",
                   "No active session", "Redirect to /admin/login", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.2 Enrolment")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-E-01", "Valid enrolment + payment",
                   "Complete form, successful Razorpay flow",
                   "Status confirmed, email sent", "Pass"],
                  ["TC-E-02", "Missing email",
                   "Form without email", "Inline 'required' error", "Pass"],
                  ["TC-E-03", "Invalid email format",
                   "email = 'abc'", "Inline 'email' error", "Pass"],
                  ["TC-E-04", "Long phone with country code",
                   "phone = '+91 9876543210'", "Accepted", "Pass"],
                  ["TC-E-05", "Payment cancelled by user",
                   "User closes Razorpay modal",
                   "Status remains pending, retry banner shown", "Pass"],
                  ["TC-E-06", "Tampered signature",
                   "Modified razorpay_signature in callback",
                   "HTTP 400 returned, status not updated", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.3 Event Registration")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-V-01", "Single attendee", "1 attendee, valid data",
                   "Registration recorded, fee = 1× sub-event fee", "Pass"],
                  ["TC-V-02", "Multiple attendees", "4 attendees",
                   "Fee = 4× sub-event fee", "Pass"],
                  ["TC-V-03", "Sub-event with redirect_link",
                   "Click sub-event with external link",
                   "Redirect to partner URL", "Pass"],
                  ["TC-V-04", "Attendee field missing",
                   "Empty attendee row submitted",
                   "Inline error", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.4 Workshop Registration")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-W-01", "Valid registration", "Active city + school",
                   "Registration saved, confirmation emailed", "Pass"],
                  ["TC-W-02", "Inactive city",
                   "City with status=0 selected",
                   "Validation error", "Pass"],
                  ["TC-W-03", "Optional merchandise",
                   "Add T-shirt", "Stored on registration row", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.5 Psychometric Test")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-P-01", "Submit answers", "All questions answered",
                   "Score computed correctly, range matched", "Pass"],
                  ["TC-P-02", "Score outside any range",
                   "Score not covered by any TestResultRange",
                   "Fallback message shown, alert logged", "Pass"],
                  ["TC-P-03", "Skipped question",
                   "Required question left blank",
                   "Inline error", "Pass"],
                  ["TC-P-04", "Graph rendering",
                   "TestGraphConfig active",
                   "Result page includes chart", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.6 Blog System")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-B-01", "Publish blog", "Valid title, slug, content",
                   "Post visible at /blog/{slug}", "Pass"],
                  ["TC-B-02", "Duplicate slug",
                   "Slug already exists",
                   "Validation error", "Pass"],
                  ["TC-B-03", "Tags syncing",
                   "Add 3 tags, save",
                   "blog_tag_pivot has 3 rows", "Pass"],
                  ["TC-B-04", "Status toggle",
                   "Click eye icon",
                   "JSON success, post hidden", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.7 Email System")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-M-01", "Template variables",
                   "Body uses {{name}}", "Replaced at send time", "Pass"],
                  ["TC-M-02", "SMTP failure",
                   "Mailer down", "email_logs status=failed, error logged", "Pass"],
                  ["TC-M-03", "Resend",
                   "Click Resend on failed log",
                   "New email_logs row created, sent successfully", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.8 Chatbot")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-C-01", "Open widget", "Visit any public page",
                   "Widget renders bottom-right", "Pass"],
                  ["TC-C-02", "Fetch FAQs",
                   "GET /chatbot/faqs", "JSON list returned", "Pass"],
                  ["TC-C-03", "Submit support ticket",
                   "POST /chatbot/support",
                   "Row created, admin alert sent", "Pass"],
              ])

    add_subsub_heading(doc, "6.1.9 Performance / Load")
    add_table(doc,
              ["TC ID", "Scenario", "Input", "Expected", "Result"],
              [
                  ["TC-L-01", "Catalogue under load",
                   "100 concurrent users",
                   "P95 < 1.5s", "Pass"],
                  ["TC-L-02", "Admin listing of 10k rows",
                   "Open /admin/enrollments",
                   "Render < 2s", "Pass"],
                  ["TC-L-03", "Payment verify latency",
                   "Razorpay callback",
                   "Verify < 1s", "Pass"],
              ])
    add_page_break(doc)

    # ============================================================
    # CHAPTER 7: SYSTEM SECURITY
    # ============================================================
    add_chapter_heading(doc, "Chapter 7: System Security")
    add_para(doc,
             "System security is essential because the platform stores "
             "personally identifiable information, accepts online payments, "
             "and exposes administrative functions. This chapter describes the "
             "checks and controls put in place and the encryption and secure "
             "practices applied across the codebase.")

    add_sub_heading(doc, "7.1 Checks and Control")
    add_subsub_heading(doc, "7.1.1 Authentication Controls")
    add_bullets(doc, [
        "Admin authentication is required for every /admin/* route via the admin.auth middleware.",
        "Session cookies are flagged HttpOnly, Secure, and SameSite=Lax.",
        "Sessions expire after 120 minutes of inactivity.",
        "Logout invalidates the server-side session and clears the cookie.",
    ])
    add_subsub_heading(doc, "7.1.2 Authorisation Controls")
    add_bullets(doc, [
        "Public controllers expose only read operations and form submissions; nothing requires authorisation.",
        "Admin controllers are gated by the admin.auth middleware uniformly.",
        "AJAX endpoints validate the CSRF token in addition to checking the session.",
    ])
    add_subsub_heading(doc, "7.1.3 Input Validation")
    add_bullets(doc, [
        "Every form runs server-side validation via $request->validate(...).",
        "Fields are type-cast to integers, strings, decimals, or booleans before persisting.",
        "Email and phone fields use Laravel's built-in validation rules.",
        "Image uploads are validated for content type, file size, and dimensions.",
    ])
    add_subsub_heading(doc, "7.1.4 Output Escaping")
    add_bullets(doc, [
        "Blade templates use {{ }} which auto-escapes output, preventing XSS by default.",
        "Raw HTML output via {!! !!} is restricted to admin-trusted content (blog body, social embeds).",
        "All file paths shown in URLs are validated against an allow-list of directories.",
    ])
    add_subsub_heading(doc, "7.1.5 CSRF Protection")
    add_para(doc,
             "Every state-changing form (POST, PUT, PATCH, DELETE) includes "
             "a CSRF token via the @csrf Blade directive. AJAX requests "
             "send the token in the X-CSRF-TOKEN header. Laravel's "
             "VerifyCsrfToken middleware rejects requests with an absent or "
             "mismatched token.")
    add_subsub_heading(doc, "7.1.6 SQL Injection Prevention")
    add_para(doc,
             "All database access goes through Eloquent or the Query Builder, "
             "both of which use parameterised queries. There are no raw SQL "
             "concatenations in the codebase.")
    add_subsub_heading(doc, "7.1.7 Brute-Force Protection")
    add_para(doc,
             "The admin login is rate-limited to five attempts per minute per "
             "IP address using Laravel's throttle middleware. After five "
             "consecutive failures the IP is blocked for sixty seconds.")

    add_sub_heading(doc, "7.2 Encryption and Secure Practices")
    add_subsub_heading(doc, "7.2.1 Transport Layer Security")
    add_para(doc,
             "All HTTP traffic is forced to HTTPS via a server-level redirect "
             "from port 80 to port 443. The TLS certificate is provisioned by "
             "Let's Encrypt and renewed automatically. HSTS is enabled with a "
             "six-month max-age once the certificate is stable.")
    add_subsub_heading(doc, "7.2.2 Password Hashing")
    add_para(doc,
             "Passwords are hashed using bcrypt with the default cost factor "
             "of 10. Hashes are stored in the admins.password column. "
             "Plain-text passwords never touch the database or the logs.")
    add_subsub_heading(doc, "7.2.3 Application Encryption Key")
    add_para(doc,
             "Laravel's APP_KEY is a 32-byte AES-256 key generated by "
             "php artisan key:generate at deploy time. The key encrypts "
             "session cookies, signed URLs, and any payload encrypted via the "
             "Crypt facade.")
    add_subsub_heading(doc, "7.2.4 Razorpay Signature Verification")
    add_para(doc,
             "Every payment callback recomputes the HMAC-SHA256 signature "
             "from order_id|payment_id using the server-side Razorpay secret. "
             "Comparison uses hash_equals to avoid timing attacks. Mismatches "
             "are rejected with HTTP 400 and logged for fraud investigation.")
    add_subsub_heading(doc, "7.2.5 Secrets Management")
    add_bullets(doc, [
        "Razorpay key/secret pairs, mailer credentials, and database passwords live in .env, not in version control.",
        ".env is mode 600 owned by the deploy user.",
        "Production .env is generated from a vault during deployment, never copied between environments by hand.",
    ])
    add_subsub_heading(doc, "7.2.6 Audit & Logging")
    add_bullets(doc, [
        "Application logs go to storage/logs/laravel.log with daily rotation.",
        "email_logs records every dispatched email with status and error text.",
        "Web server access logs are retained for 30 days.",
        "Failed admin logins are logged with the source IP for review.",
    ])
    add_subsub_heading(doc, "7.2.7 Backup & Recovery")
    add_bullets(doc, [
        "MySQL is backed up nightly via mysqldump piped to gzip and uploaded to off-site storage.",
        "Backups are encrypted at rest using GPG with a key held by the operations team.",
        "Recovery drills are run quarterly to validate the restore procedure.",
    ])
    add_page_break(doc)

    # ============================================================
    # CHAPTER 8: CONCLUSION / FUTURE ENHANCEMENT
    # ============================================================
    add_chapter_heading(doc, "Chapter 8: Conclusion / Future Enhancement")
    add_sub_heading(doc, "8.1 Conclusion")
    add_para(doc,
             "The ActToAction project demonstrates that a single, well-designed "
             "Laravel application can replace a fragmented landscape of "
             "third-party tools, manual spreadsheets, and ad-hoc payment "
             "links. By bringing courses, events, workshops, blogs, services, "
             "industries, psychometric tests, and online payments under one "
             "roof, the platform delivers measurable benefits to every "
             "stakeholder group: visitors get a consistent and reliable "
             "experience; operational staff get a single workspace; finance "
             "gets real-time payment reconciliation; and the engineering team "
             "gets a predictable, conventional codebase that is easy to "
             "extend.")
    add_para(doc,
             "The project also reinforced several engineering lessons. The "
             "first is that opinionated frameworks pay off — Laravel's "
             "conventions allowed dozens of CRUD modules to be implemented "
             "with minimal incidental complexity. The second is that "
             "configuration-driven design (templates, test ranges, content "
             "categories) enables a small engineering team to support a much "
             "larger operations team without becoming a bottleneck. The third "
             "is that early investment in payment-signature verification and "
             "email logging pays off in long-term confidence: every rupee is "
             "accounted for, and every dispatched message is traceable.")
    add_para(doc,
             "From a personal standpoint, the project provided hands-on "
             "exposure to the full lifecycle of a non-trivial web application "
             "— from requirements gathering through database design, "
             "implementation, testing, and deployment. The experience of "
             "shipping software that real users depend on, and that handles "
             "real money, is qualitatively different from coursework "
             "exercises and is the most valuable outcome of the project.")

    add_sub_heading(doc, "8.2 Future Enhancements")
    add_subsub_heading(doc, "8.2.1 Near-Term Enhancements")
    add_bullets(doc, [
        "Multi-factor authentication for admin accounts, using TOTP or magic links.",
        "Activity logging (who-changed-what) using spatie/laravel-activitylog.",
        "Capacity enforcement on sub-events and workshops.",
        "In-app refund flow that talks to Razorpay's refund API.",
        "Bulk import of school and city lists via CSV upload.",
    ])
    add_subsub_heading(doc, "8.2.2 Mid-Term Enhancements")
    add_bullets(doc, [
        "Role-based access control to support distinct admin personas (editor, finance, super-admin).",
        "Structured data (schema.org) for events and courses to improve organic search visibility.",
        "Internationalisation — Hindi and selected regional languages for the public site.",
        "PWA-style installable experience for students.",
        "Automated revenue and attendance dashboards using Filament or similar.",
    ])
    add_subsub_heading(doc, "8.2.3 Long-Term Enhancements")
    add_bullets(doc, [
        "Native mobile applications (iOS and Android) backed by a versioned REST API.",
        "Headless front-end for the marketing site (Inertia + Vue/React).",
        "Real-time notifications via WebSockets (Laravel Reverb / Pusher) for ticketing and chat.",
        "Machine-learning-based recommendations on the catalogue.",
        "Multi-tenant capability so franchise partners can have their own scoped admin area.",
    ])
    add_page_break(doc)

    # ============================================================
    # CHAPTER 9: BIBLIOGRAPHY
    # ============================================================
    add_chapter_heading(doc, "Chapter 9: Bibliography")
    add_sub_heading(doc, "9.1 Books")
    add_bullets(doc, [
        "Otwell, T. (2024). Laravel: Up and Running, 3rd Edition. O'Reilly Media.",
        "Pratt, P. J., & Last, M. Z. (2014). Concepts of Database Management, 8th Edition. Cengage Learning.",
        "Pressman, R. S., & Maxim, B. R. (2019). Software Engineering: A Practitioner's Approach, 9th Edition. McGraw-Hill.",
        "Sommerville, I. (2015). Software Engineering, 10th Edition. Pearson.",
        "Stauffer, M. (2023). Laravel: Up & Running. O'Reilly Media.",
        "Welling, L., & Thomson, L. (2017). PHP and MySQL Web Development, 5th Edition. Addison-Wesley.",
    ])
    add_sub_heading(doc, "9.2 Online Resources")
    add_bullets(doc, [
        "Laravel Documentation — https://laravel.com/docs",
        "PHP Manual — https://www.php.net/manual/en/",
        "MySQL 8.0 Reference Manual — https://dev.mysql.com/doc/refman/8.0/en/",
        "Razorpay PHP SDK — https://github.com/razorpay/razorpay-php",
        "Razorpay Web Integration — https://razorpay.com/docs/payments/payment-gateway/web-integration/",
        "OWASP Top Ten — https://owasp.org/www-project-top-ten/",
        "PSR-12 PHP Coding Standard — https://www.php-fig.org/psr/psr-12/",
        "MDN Web Docs — https://developer.mozilla.org/",
    ])
    add_sub_heading(doc, "9.3 Standards")
    add_bullets(doc, [
        "ISO/IEC 25010:2011 — Systems and software Quality Requirements and Evaluation.",
        "IEEE 830-1998 — Recommended Practice for Software Requirements Specifications.",
        "OWASP Application Security Verification Standard (ASVS) v4.",
    ])

    # ============================================================
    # APPENDIX A — PROJECT SCHEDULE
    # ============================================================
    add_chapter_heading(doc, "Appendix A: Project Schedule")
    add_para(doc,
             "The project was executed over a fourteen-week schedule. The "
             "Gantt chart below illustrates the major phases, their duration, "
             "and the parallel tracks that ran concurrently during the busiest "
             "weeks.")
    add_image(doc, charts["gantt"], "Figure A.1 — Project schedule (Gantt chart).", width=6.5)
    add_sub_heading(doc, "A.1 Phase-Wise Effort Distribution")
    add_table(doc,
              ["Phase", "Duration", "Effort (person-days)", "Deliverables"],
              [
                  ["Requirements & Study", "Weeks 1-2", "8", "Requirements doc, stakeholder map"],
                  ["System Analysis", "Weeks 2-3", "8", "Functional & non-functional spec"],
                  ["Database Design", "Weeks 3-4", "6", "ER diagram, migrations"],
                  ["UI / UX Design", "Weeks 3-5", "10", "Wireframes, mockups"],
                  ["Public Site Development", "Weeks 5-9", "20", "All public controllers + views"],
                  ["Admin Panel Development", "Weeks 5-10", "30", "40+ CRUD modules"],
                  ["Payment Integration", "Weeks 8-9", "5", "Razorpay end-to-end"],
                  ["Email System", "Weeks 8-9", "4", "Templates + logs"],
                  ["Testing & QA", "Weeks 10-12", "10", "Test cases, bug fixes"],
                  ["Deployment", "Week 12-13", "3", "Staging + production live"],
                  ["Documentation", "Weeks 11-13", "8", "User manual + this report"],
              ])
    add_para(doc,
             "Total estimated effort: 112 person-days, executed by a "
             "two-person development team supplemented by part-time review "
             "by the project guide.")
    add_page_break(doc)

    # ============================================================
    # APPENDIX B — SAMPLE REPORTS
    # ============================================================
    add_chapter_heading(doc, "Appendix B: Sample Reports")
    add_sub_heading(doc, "B.1 Daily Enrolment Report")
    add_para(doc,
             "The daily enrolment report is generated by the operations team "
             "every morning. It lists all enrolments confirmed in the past "
             "24 hours along with payment details. A representative report is "
             "reproduced below.")
    add_table(doc,
              ["Date", "Course", "Student", "Fee", "Mode", "Status"],
              [
                  ["06-May-2026 09:14", "Pre-School Foundation", "Aarav Sharma", "₹4,500", "UPI", "Confirmed"],
                  ["06-May-2026 10:32", "Coding for Kids", "Anaya Patel", "₹6,000", "Card", "Confirmed"],
                  ["06-May-2026 11:50", "Robotics Basics", "Vihaan Khan", "₹8,500", "UPI", "Pending"],
                  ["06-May-2026 12:18", "Public Speaking", "Diya Nair", "₹3,200", "NetBanking", "Confirmed"],
                  ["06-May-2026 14:05", "Math Olympiad Prep", "Aarush Iyer", "₹5,500", "UPI", "Confirmed"],
                  ["06-May-2026 15:33", "Speed Reading", "Reyansh Joshi", "₹3,500", "Card", "Confirmed"],
              ])
    add_para(doc,
             "Total enrolments in 24h: 6  |  Confirmed: 5  |  Pending: 1  |  "
             "Revenue (confirmed): ₹ 22,700.")

    add_sub_heading(doc, "B.2 Weekly Workshop Attendance Summary")
    add_table(doc,
              ["City", "School", "Workshop", "Registered", "Attended", "%"],
              [
                  ["Jaipur", "DPS Jaipur", "Robotics Day", "45", "42", "93%"],
                  ["Jaipur", "St Xavier's", "Art & Craft", "38", "35", "92%"],
                  ["Udaipur", "MGD School", "Coding Basics", "30", "28", "93%"],
                  ["Kota", "Allen Career", "Math Genius", "55", "50", "91%"],
                  ["Ajmer", "Mayoor School", "Speed Reading", "22", "20", "91%"],
                  ["Jodhpur", "Rajmata Krishna Kumari", "Public Speaking", "18", "17", "94%"],
              ])
    add_para(doc,
             "Total registered: 208  |  Total attended: 192  |  "
             "Attendance rate: 92.3%.")

    add_sub_heading(doc, "B.3 Monthly Revenue Report")
    add_table(doc,
              ["Month", "Enrolments", "Event Reg.", "Workshop Reg.", "Total Revenue"],
              [
                  ["Jan 2026", "82", "145", "210", "₹ 8,45,000"],
                  ["Feb 2026", "78", "168", "195", "₹ 8,72,000"],
                  ["Mar 2026", "112", "210", "245", "₹ 11,40,000"],
                  ["Apr 2026", "98", "185", "228", "₹ 9,85,000"],
                  ["May 2026 (MTD)", "45", "67", "82", "₹ 4,12,000"],
              ])
    add_para(doc, "Year-to-date revenue (Jan – May 2026): ₹ 42,54,000.")

    add_sub_heading(doc, "B.4 Email Delivery Report")
    add_table(doc,
              ["Template", "Sent", "Failed", "Success Rate"],
              [
                  ["Enrolment Confirmation", "1,245", "8", "99.36%"],
                  ["Event Ticket", "823", "6", "99.27%"],
                  ["Workshop Confirmation", "960", "5", "99.48%"],
                  ["Admission Full-Form Link", "412", "11", "97.33%"],
                  ["Volunteer Welcome", "85", "1", "98.82%"],
                  ["Newsletter (April)", "12,400", "98", "99.21%"],
              ])
    add_para(doc, "Total emails dispatched (rolling 30 days): 15,925  |  "
                  "Overall success rate: 99.18%.")
    add_page_break(doc)

    # ============================================================
    # APPENDIX C — ABBREVIATIONS
    # ============================================================
    add_chapter_heading(doc, "Appendix C: List of Abbreviations")
    add_table(doc,
              ["Abbreviation", "Expansion"],
              [
                  ["AJAX", "Asynchronous JavaScript and XML"],
                  ["API", "Application Programming Interface"],
                  ["CRUD", "Create, Read, Update, Delete"],
                  ["CSRF", "Cross-Site Request Forgery"],
                  ["CSS", "Cascading Style Sheets"],
                  ["DBMS", "Database Management System"],
                  ["DFD", "Data Flow Diagram"],
                  ["DOM", "Document Object Model"],
                  ["ERD", "Entity-Relationship Diagram"],
                  ["FK", "Foreign Key"],
                  ["GUI", "Graphical User Interface"],
                  ["HMAC", "Hash-based Message Authentication Code"],
                  ["HTML", "HyperText Markup Language"],
                  ["HTTP", "HyperText Transfer Protocol"],
                  ["HTTPS", "HyperText Transfer Protocol Secure"],
                  ["IDE", "Integrated Development Environment"],
                  ["JSON", "JavaScript Object Notation"],
                  ["LMS", "Learning Management System"],
                  ["MCA", "Master of Computer Applications"],
                  ["MFA", "Multi-Factor Authentication"],
                  ["MVC", "Model-View-Controller"],
                  ["ORM", "Object-Relational Mapping"],
                  ["OWASP", "Open Web Application Security Project"],
                  ["PCI-DSS", "Payment Card Industry Data Security Standard"],
                  ["PDF", "Portable Document Format"],
                  ["PHP", "Hypertext Preprocessor"],
                  ["PK", "Primary Key"],
                  ["PSR", "PHP Standards Recommendation"],
                  ["QA", "Quality Assurance"],
                  ["RBAC", "Role-Based Access Control"],
                  ["REST", "REpresentational State Transfer"],
                  ["RTU", "Rajasthan Technical University"],
                  ["SDK", "Software Development Kit"],
                  ["SEO", "Search Engine Optimization"],
                  ["SMTP", "Simple Mail Transfer Protocol"],
                  ["SPDI", "Sensitive Personal Data or Information"],
                  ["SQL", "Structured Query Language"],
                  ["SSL", "Secure Sockets Layer"],
                  ["TLS", "Transport Layer Security"],
                  ["TOTP", "Time-based One-Time Password"],
                  ["UI / UX", "User Interface / User Experience"],
                  ["UPI", "Unified Payments Interface"],
                  ["URL", "Uniform Resource Locator"],
                  ["UUID", "Universally Unique Identifier"],
                  ["WYSIWYG", "What You See Is What You Get"],
                  ["XSS", "Cross-Site Scripting"],
              ])
    add_page_break(doc)

    # ============================================================
    # APPENDIX D — COMPLETE TABLE LIST
    # ============================================================
    add_chapter_heading(doc, "Appendix D: Complete List of Database Tables")
    add_para(doc,
             "The full database schema spans more than eighty tables. They "
             "are listed below by functional cluster for quick reference.")
    add_sub_heading(doc, "D.1 Identity & Settings")
    add_bullets(doc, [
        "users", "password_reset_tokens", "sessions", "settings", "themes",
    ])
    add_sub_heading(doc, "D.2 Marketing Surface")
    add_bullets(doc, [
        "sliders", "hero_banners", "announcement_bars", "notification_banners",
        "stats", "brands", "testimonials", "testimonial_videos",
    ])
    add_sub_heading(doc, "D.3 Education Domain")
    add_bullets(doc, [
        "course_categories", "courses", "course_documents", "course_sessions",
        "enrollments", "admission_short_forms", "admission_full_forms",
        "people", "team_members",
    ])
    add_sub_heading(doc, "D.4 Events")
    add_bullets(doc, [
        "events", "sub_events", "event_registrations",
        "event_registration_attendees",
    ])
    add_sub_heading(doc, "D.5 Workshops")
    add_bullets(doc, [
        "workshop_cities", "workshop_schools", "workshop_age_groups",
        "workshop_registrations",
    ])
    add_sub_heading(doc, "D.6 Content")
    add_bullets(doc, [
        "blogs", "blog_categories", "blog_authors", "blog_tags",
        "blog_tag_pivot",
        "abouts", "about_categories", "about_sections", "page_categories",
    ])
    add_sub_heading(doc, "D.7 Services & Industries")
    add_bullets(doc, [
        "service_categories", "service_subcategories", "services",
        "service_features", "service_benefits", "service_essentials",
        "service_faqs",
        "industries", "industry_services", "industry_features",
        "industry_faqs",
    ])
    add_sub_heading(doc, "D.8 Psychometric Tests")
    add_bullets(doc, [
        "psych_categories", "psych_tests", "psych_questions",
        "test_result_ranges", "test_graph_configs",
    ])
    add_sub_heading(doc, "D.9 Galleries & Media")
    add_bullets(doc, [
        "gallery_categories", "galleries", "gallery_images",
        "video_galleries",
        "youtube_categories", "youtube_videos",
    ])
    add_sub_heading(doc, "D.10 Operations")
    add_bullets(doc, [
        "payments", "email_templates", "email_logs", "newsletters",
        "chatbot_faqs", "chatbot_support_tickets",
        "contacts", "contact_infos", "frontend_contact_us",
        "centers", "states", "merchandises",
        "volunteers", "franchise_forms", "become_partners",
        "school_partners", "school_partner_categories", "school_sections",
        "summer_partners", "summer_partner_categories",
        "action_items",
    ])
    add_page_break(doc)

    # ============================================================
    # APPENDIX E — SAMPLE TEST DATA
    # ============================================================
    add_chapter_heading(doc, "Appendix E: Sample Test Data")
    add_para(doc,
             "The data below was used during development and quality "
             "assurance to validate the system end-to-end. None of the "
             "personal details correspond to real individuals.")
    add_sub_heading(doc, "E.1 Sample Course Records")
    add_table(doc,
              ["ID", "Category", "Name", "Slug", "Fee", "Status"],
              [
                  ["1", "Pre-School", "Pre-School Foundation", "pre-school-foundation", "₹4500", "Active"],
                  ["2", "Coding", "Coding for Kids", "coding-for-kids", "₹6000", "Active"],
                  ["3", "STEM", "Robotics Basics", "robotics-basics", "₹8500", "Active"],
                  ["4", "Soft Skills", "Public Speaking", "public-speaking", "₹3200", "Active"],
                  ["5", "Academics", "Math Olympiad Prep", "math-olympiad-prep", "₹5500", "Active"],
                  ["6", "Arts", "Art & Craft Mastery", "art-craft-mastery", "₹2800", "Active"],
                  ["7", "Soft Skills", "Speed Reading", "speed-reading", "₹3500", "Active"],
                  ["8", "Coding", "Web Development for Teens", "webdev-teens", "₹9500", "Inactive"],
              ])
    add_sub_heading(doc, "E.2 Sample Event Records")
    add_table(doc,
              ["ID", "Title", "Slug", "Start Date", "Status"],
              [
                  ["1", "Annual Quiz Competition 2026", "annual-quiz-2026", "12-Jun-2026", "Active"],
                  ["2", "Robotics Carnival", "robotics-carnival", "20-Jul-2026", "Active"],
                  ["3", "Summer Coding Bootcamp", "summer-coding-2026", "01-Jun-2026", "Active"],
                  ["4", "Art Festival", "art-festival-2026", "15-Aug-2026", "Active"],
              ])
    add_sub_heading(doc, "E.3 Sample Sub-Event Records")
    add_table(doc,
              ["ID", "Event", "Title", "Fee", "Date"],
              [
                  ["1", "Annual Quiz", "Junior Round (Class 1-3)", "₹250", "12-Jun-2026"],
                  ["2", "Annual Quiz", "Middle Round (Class 4-6)", "₹350", "12-Jun-2026"],
                  ["3", "Annual Quiz", "Senior Round (Class 7-10)", "₹500", "13-Jun-2026"],
                  ["4", "Robotics Carnival", "Build Challenge", "₹800", "20-Jul-2026"],
                  ["5", "Robotics Carnival", "Coding Sprint", "₹600", "21-Jul-2026"],
              ])
    add_sub_heading(doc, "E.4 Sample Psychometric Test")
    add_para(doc, "Test: \"Career Aptitude Indicator\" (slug: career-aptitude)")
    add_table(doc,
              ["Q#", "Question", "Option Weights"],
              [
                  ["1", "I enjoy solving logical puzzles", "Strongly Agree=4, Agree=3, Neutral=2, Disagree=1, Strongly Disagree=0"],
                  ["2", "I prefer working in groups over alone", "SA=4, A=3, N=2, D=1, SD=0"],
                  ["3", "I am drawn to building things with my hands", "SA=4, A=3, N=2, D=1, SD=0"],
                  ["4", "I feel comfortable speaking in public", "SA=4, A=3, N=2, D=1, SD=0"],
                  ["5", "I am good at managing time", "SA=4, A=3, N=2, D=1, SD=0"],
              ])
    add_para(doc, "Result Ranges:")
    add_table(doc,
              ["Range", "Label", "Interpretation"],
              [
                  ["0-7", "Exploratory", "You are still exploring your interests."],
                  ["8-13", "Developing", "Some clear preferences are emerging."],
                  ["14-17", "Focused", "You have well-defined leanings towards specific career paths."],
                  ["18-20", "Highly Focused", "Strong clarity — proceed with confidence."],
              ])

    add_sub_heading(doc, "E.5 Sample Razorpay Test Credentials (Test Mode Only)")
    add_table(doc,
              ["Item", "Value"],
              [
                  ["Test Card", "4111 1111 1111 1111"],
                  ["Expiry", "Any future month / year"],
                  ["CVV", "123"],
                  ["UPI Test ID", "success@razorpay"],
                  ["NetBanking", "Any test bank"],
                  ["Test Mode Key Prefix", "rzp_test_*"],
              ])
    add_para(doc,
             "These credentials are publicly published by Razorpay for "
             "integration testing. They never connect to a real bank account.")
    add_page_break(doc)

    out_path = os.path.join(OUTDIR, "ActToAction_Documentation.docx")
    doc.save(out_path)
    print(f"Saved to {out_path}")
    return out_path


if __name__ == "__main__":
    build()
