"""
ActToAction — E-Commerce Platform Proposal & Cost Estimation
Generates a Word document covering features, requirements, architecture,
timeline, and Indian-market cost breakdown.
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(OUTDIR, "proposal_images")
REAL_IMG_DIR = os.path.join(OUTDIR, "real_images")
os.makedirs(IMG_DIR, exist_ok=True)

FONT = "Times New Roman"
NAVY = RGBColor(0x1F, 0x38, 0x64)
DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x60, 0x60, 0x60)
ACCENT = RGBColor(0xC0, 0x50, 0x4D)


# ----------------------- diagram helpers -----------------------

def _box(ax, x, y, w, h, text, color="#4F81BD", text_color="white", fontsize=9):
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle="round,pad=0.02,rounding_size=0.08",
                         linewidth=1.2, facecolor=color, edgecolor="#1F3864")
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fontsize, color=text_color, weight="bold", wrap=True)


def _arrow(ax, x1, y1, x2, y2, label="", color="#404040"):
    arr = FancyArrowPatch((x1, y1), (x2, y2),
                          arrowstyle="-|>", mutation_scale=14,
                          color=color, linewidth=1.3)
    ax.add_patch(arr)
    if label:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2, label, fontsize=8,
                color="#202060", ha="center", va="center",
                bbox=dict(facecolor="white", edgecolor="none", pad=1))


def _new(w=10, h=6, title=""):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(0, w)
    ax.set_ylim(0, h)
    ax.set_aspect("equal")
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=12, weight="bold", pad=10)
    return fig, ax


def save(name):
    p = os.path.join(IMG_DIR, name + ".png")
    plt.tight_layout()
    plt.savefig(p, dpi=170, bbox_inches="tight", facecolor="white")
    plt.close()
    return p


def chart_architecture():
    fig, ax = _new(11, 7.5, "Proposed System Architecture — ActToAction E-Commerce")
    _box(ax, 0.3, 6.4, 2.2, 0.7, "Customer (Web)", "#7EB6E1")
    _box(ax, 3.0, 6.4, 2.2, 0.7, "Customer (Mobile)", "#7EB6E1")
    _box(ax, 5.7, 6.4, 2.2, 0.7, "Seller / Vendor", "#9BBB59")
    _box(ax, 8.4, 6.4, 2.4, 0.7, "Admin / Operations", "#A8D08D")

    _box(ax, 0.5, 5.2, 10.0, 0.7, "CDN  +  WAF  +  HTTPS", "#305496")
    _box(ax, 0.5, 4.0, 10.0, 0.7, "Laravel Application — Routing, Controllers, Middleware", "#1F3864")

    services = [
        ("Catalog", 0.3, 2.8), ("Cart / Wishlist", 1.8, 2.8),
        ("Orders", 3.3, 2.8), ("Payments", 4.8, 2.8),
        ("Returns", 6.3, 2.8), ("Coupons", 7.8, 2.8),
        ("BOPIS", 9.3, 2.8),
    ]
    for n, x, y in services:
        _box(ax, x, y, 1.4, 0.7, n, "#4F81BD", fontsize=8)
    services2 = [
        ("i18n / FX", 0.3, 1.9), ("Chatbot", 1.8, 1.9),
        ("Push", 3.3, 1.9), ("Email", 4.8, 1.9),
        ("Gift Builder", 6.3, 1.9), ("Analytics", 7.8, 1.9),
        ("Spin Wheel", 9.3, 1.9),
    ]
    for n, x, y in services2:
        _box(ax, x, y, 1.4, 0.7, n, "#4F81BD", fontsize=8)

    _box(ax, 0.3, 0.9, 2.5, 0.7, "MySQL  (primary)", "#1F3864")
    _box(ax, 3.0, 0.9, 2.5, 0.7, "Redis (cache / queue)", "#1F3864")
    _box(ax, 5.7, 0.9, 2.5, 0.7, "S3-compatible Storage", "#1F3864")
    _box(ax, 8.4, 0.9, 2.4, 0.7, "Elasticsearch (search)", "#1F3864")

    _box(ax, 0.3, 0.0, 2.5, 0.6, "Razorpay / Stripe", "#C0504D", fontsize=8)
    _box(ax, 3.0, 0.0, 2.5, 0.6, "DHL / BlueDart / Delhivery", "#C0504D", fontsize=8)
    _box(ax, 5.7, 0.0, 2.5, 0.6, "FCM / OneSignal Push", "#C0504D", fontsize=8)
    _box(ax, 8.4, 0.0, 2.4, 0.6, "SMTP / SendGrid", "#C0504D", fontsize=8)

    for x in [1.4, 4.1, 6.8, 9.6]:
        _arrow(ax, x, 6.4, x, 5.9)
        _arrow(ax, x, 5.2, x, 4.7)
        _arrow(ax, x, 4.0, x, 3.5)
    return save("01_architecture")


def chart_user_flow():
    fig, ax = _new(11, 6.5, "Customer Purchase Flow")
    steps = [
        ("Land on Site", "#4F81BD"),
        ("Browse Catalog", "#4F81BD"),
        ("Filter / Search", "#9BBB59"),
        ("Add to Cart / Wishlist", "#9BBB59"),
        ("Apply Coupon / Spin Wheel", "#F4B084"),
        ("Choose Delivery or BOPIS", "#9BBB59"),
        ("Guest or Sign-in Checkout", "#305496"),
        ("Pay (Razorpay / Stripe)", "#C0504D"),
        ("Order Confirmation + Invoice", "#9BBB59"),
        ("Track Shipment (DHL etc.)", "#4F81BD"),
        ("Receive Goods", "#4F81BD"),
        ("Optional Return / Feedback", "#F4B084"),
    ]
    cols = 4
    for i, (lbl, c) in enumerate(steps):
        col = i % cols
        row = i // cols
        x = 0.3 + col * 2.7
        y = 5.5 - row * 1.8
        _box(ax, x, y, 2.4, 0.9, lbl, c)
        if i > 0:
            pcol = (i - 1) % cols
            prow = (i - 1) // cols
            px = 0.3 + pcol * 2.7
            py = 5.5 - prow * 1.8
            if prow == row:
                _arrow(ax, px + 2.4, py + 0.45, x, y + 0.45)
            else:
                _arrow(ax, px + 1.2, py, x + 1.2, y + 0.9)
    return save("02_user_flow")


def chart_gantt():
    fig, ax = _new(11, 6, "Project Schedule — 22 Weeks")
    phases = [
        ("Discovery & Wireframes", 0, 3, "#4F81BD"),
        ("UI/UX Design", 2, 4, "#4F81BD"),
        ("Backend Foundation", 4, 3, "#305496"),
        ("Catalog + Cart", 6, 4, "#9BBB59"),
        ("Payments + Orders", 9, 3, "#C0504D"),
        ("i18n + Currency", 9, 2, "#9BBB59"),
        ("Shipping (DHL etc.)", 11, 2, "#9BBB59"),
        ("BOPIS (Store Pickup)", 12, 2, "#F4B084"),
        ("Returns + Invoicing", 12, 2, "#9BBB59"),
        ("Coupons + Spin Wheel", 13, 2, "#F4B084"),
        ("Gift Builder + Notes", 14, 2, "#F4B084"),
        ("Chatbot + Push + Email", 14, 3, "#9BBB59"),
        ("Admin Panel", 4, 12, "#A8D08D"),
        ("QA + UAT", 17, 3, "#9BBB59"),
        ("Deployment + Handover", 20, 2, "#4F81BD"),
    ]
    weeks = 22
    ax.set_xlim(-3.5, weeks + 0.5)
    ax.set_ylim(-0.5, len(phases) + 0.5)
    ax.set_aspect("auto")
    for i in range(weeks + 1):
        ax.plot([i, i], [-0.5, len(phases) + 0.5],
                color="#E0E0E0", linewidth=0.5)
        if i % 2 == 0:
            ax.text(i + 0.5, len(phases) + 0.2, f"W{i+1}",
                    fontsize=7, ha="center", color="#606060")
    for i, (n, s, d, c) in enumerate(phases):
        y = len(phases) - i - 1
        ax.add_patch(Rectangle((s, y - 0.3), d, 0.6,
                     facecolor=c, edgecolor="#202020", linewidth=0.5))
        ax.text(-0.3, y, n, fontsize=9, ha="right", va="center",
                weight="bold", color="#1F3864")
        ax.text(s + d / 2, y, f"{d}w", fontsize=8, ha="center",
                va="center", color="white", weight="bold")
    ax.set_xticks([])
    ax.set_yticks([])
    return save("03_gantt")


def chart_cost_breakdown():
    fig, ax = plt.subplots(figsize=(10, 6))
    labels = ["Design (UI/UX)", "Backend Dev", "Frontend Dev",
             "Integrations", "Admin Panel", "QA & Testing",
             "Project Mgmt", "Deployment", "Contingency"]
    values = [150000, 600000, 400000, 250000, 200000,
              150000, 120000, 50000, 180000]
    colors = ["#4F81BD", "#1F3864", "#9BBB59", "#C0504D",
              "#A8D08D", "#F4B084", "#305496", "#7EB6E1", "#9C5700"]
    explode = [0.02] * len(labels)
    ax.pie(values, labels=labels, colors=colors, explode=explode,
           autopct=lambda p: f"₹{int(p * sum(values) / 100):,}",
           startangle=90, textprops={"fontsize": 8.5, "weight": "bold"})
    ax.set_title("Mid-Tier Cost Breakdown (₹21 Lakhs total)",
                 fontsize=12, weight="bold", pad=10)
    return save("04_cost_pie")


def chart_tier_comparison():
    fig, ax = plt.subplots(figsize=(10, 5))
    tiers = ["Lean (MVP)", "Mid-Tier (Recommended)", "Premium (Enterprise)"]
    low = [800000, 1500000, 2500000]
    high = [1200000, 2200000, 4000000]
    x = range(len(tiers))
    ax.bar([i - 0.2 for i in x], low, width=0.4,
           label="Low Estimate", color="#4F81BD")
    ax.bar([i + 0.2 for i in x], high, width=0.4,
           label="High Estimate", color="#C0504D")
    for i, (lo, hi) in enumerate(zip(low, high)):
        ax.text(i - 0.2, lo + 30000, f"₹{lo/100000:.1f}L",
                ha="center", fontsize=9, weight="bold")
        ax.text(i + 0.2, hi + 30000, f"₹{hi/100000:.1f}L",
                ha="center", fontsize=9, weight="bold")
    ax.set_xticks(list(x))
    ax.set_xticklabels(tiers, fontsize=10, weight="bold")
    ax.set_ylabel("Total cost (INR)", fontsize=10)
    ax.set_title("Cost by Engagement Tier (India market)",
                 fontsize=12, weight="bold")
    ax.legend(loc="upper left")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.yaxis.set_major_formatter(
        plt.FuncFormatter(lambda x, p: f"₹{int(x/100000)}L"))
    plt.tight_layout()
    p = os.path.join(IMG_DIR, "05_tier.png")
    plt.savefig(p, dpi=170, bbox_inches="tight", facecolor="white")
    plt.close()
    return p


def chart_feature_complexity():
    fig, ax = plt.subplots(figsize=(10, 7))
    features = [
        "BOPIS (Store Pickup)", "Multilingual", "DHL Integration",
        "Suggestion Chatbot", "Spin the Wheel", "Gift Package Builder",
        "Return Management", "Multi-Currency", "Order Tracking",
        "Push Notifications", "Coupon System", "Wishlist & Favorites",
        "Cart Reminder", "Personalisation", "Custom Notes",
        "Catalog Mgmt", "Cart", "Guest Checkout",
    ]
    hours = [80, 60, 60, 60, 30, 50, 50, 40, 30, 50,
             30, 25, 30, 30, 20, 60, 35, 20]
    ax.barh(features, hours, color="#4F81BD", edgecolor="#1F3864")
    for i, h in enumerate(hours):
        ax.text(h + 0.5, i, f"{h}h  (~₹{h*900:,})",
                va="center", fontsize=8, weight="bold", color="#1F3864")
    ax.set_xlabel("Development Effort (hours)", fontsize=10)
    ax.set_title("Per-Feature Effort Estimate (rate ₹900/hr)",
                 fontsize=12, weight="bold")
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    p = os.path.join(IMG_DIR, "06_feature_effort.png")
    plt.savefig(p, dpi=170, bbox_inches="tight", facecolor="white")
    plt.close()
    return p


# ----------------------- doc helpers -----------------------

def set_font(run, size=12, bold=False, italic=False, color=None, name=FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_p(p, line_spacing=1.5, sb=0, sa=6, align=None, indent=None):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    if align is not None:
        p.alignment = align
    if indent is not None:
        pf.first_line_indent = Inches(indent)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    set_font(r, size=18, bold=True, color=NAVY)
    set_p(p, sb=12, sa=14, align=WD_ALIGN_PARAGRAPH.CENTER)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=14, bold=True, color=NAVY)
    set_p(p, sb=10, sa=6)


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=12, bold=True, color=DARK)
    set_p(p, sb=6, sa=4)


def para(doc, text, indent=0.3, justify=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=12)
    set_p(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else None,
          indent=indent)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_font(r, size=12)
        set_p(p, sa=2)


def numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(it)
        set_font(r, size=12)
        set_p(p, sa=2)


def img(doc, path, caption=None, width=6.0):
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        set_font(r, size=11, italic=True, color=GREY)
        set_p(cp, sa=10)


def pg(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def table(doc, headers, rows, header_fill="1F3864"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_fill)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            pcell = cells[i].paragraphs[0]
            rr = pcell.add_run(v)
            set_font(rr, size=10)
    p = doc.add_paragraph()
    set_p(p)


def set_default(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)


# ----------------------- build doc -----------------------

def build():
    print("Generating diagrams...")
    arch = chart_architecture()
    flow = chart_user_flow()
    gantt = chart_gantt()
    pie = chart_cost_breakdown()
    tiers = chart_tier_comparison()
    feff = chart_feature_complexity()

    print("Building document...")
    doc = Document()
    set_default(doc)
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.2)
        s.right_margin = Inches(1.0)
        s.page_height = Inches(11.69)
        s.page_width = Inches(8.27)

    # ---------- COVER ----------
    logo = os.path.join(REAL_IMG_DIR, "logo.png")
    if os.path.exists(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p(p, sb=80, sa=12)
        p.add_run().add_picture(logo, width=Inches(1.6))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_p(p, sb=12, sa=8)
    r = p.add_run("ACTTOACTION")
    set_font(r, size=36, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_p(p, sa=6)
    r = p.add_run("E-COMMERCE PLATFORM")
    set_font(r, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_p(p, sa=24)
    r = p.add_run("Project Proposal, Requirements & Cost Estimation")
    set_font(r, size=14, italic=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_p(p, sb=60, sa=6)
    r = p.add_run("Multilingual  •  Multi-Currency  •  BOPIS  •  31 Features")
    set_font(r, size=13, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_p(p, sb=80, sa=4)
    r = p.add_run("Prepared for: ActToAction Management")
    set_font(r, size=12, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_p(p, sa=4)
    r = p.add_run("Indicative Budget (India): ₹ 8 L  to  ₹ 40 L")
    set_font(r, size=13, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_p(p, sa=4)
    r = p.add_run("Version 1.0  •  May 2026")
    set_font(r, size=11, color=GREY)
    pg(doc)

    # ---------- TOC ----------
    h1(doc, "Table of Contents")
    toc = [
        "1. Executive Summary",
        "2. Project Vision & Scope",
        "3. Complete Feature List (31 modules)",
        "4. Detailed Feature Specifications",
        "5. Functional Requirements",
        "6. Non-Functional Requirements",
        "7. Technology Stack",
        "8. Proposed System Architecture",
        "9. Customer Purchase Flow",
        "10. Database Modules",
        "11. Third-Party Integrations & SaaS Costs",
        "12. Project Schedule (22-Week Gantt)",
        "13. Per-Feature Effort Estimate",
        "14. India-Market Cost Breakdown",
        "15. Three Engagement Tiers (Lean / Mid / Premium)",
        "16. Recurring Annual Costs",
        "17. Payment Milestones",
        "18. Assumptions, Inclusions & Exclusions",
        "19. Risks & Mitigation",
        "20. Post-Launch Support (AMC)",
        "21. Acceptance & Sign-Off",
    ]
    for t in toc:
        p = doc.add_paragraph()
        r = p.add_run(t)
        set_font(r, size=12)
        set_p(p, sa=2)
    pg(doc)

    # ---------- 1. EXEC SUMMARY ----------
    h1(doc, "1. Executive Summary")
    para(doc,
         "This document proposes the design, development, and launch of "
         "the ActToAction e-commerce platform — a multilingual, multi-"
         "currency online store with a unique 'Go Offline / Get It From "
         "Store' (BOPIS) capability, 30+ supporting features, and a "
         "complete back-office for catalog, order, return, and customer "
         "management.")
    para(doc,
         "The proposal covers the full feature inventory you requested, "
         "translates each item into concrete functional and non-functional "
         "requirements, recommends a technology stack and architecture, "
         "lays out a 22-week build schedule, and gives an India-market "
         "cost estimate broken down by feature, by phase, and by "
         "engagement tier.")
    h2(doc, "1.1 Headline Numbers")
    table(doc,
          ["Engagement Tier", "Scope", "Total Cost (INR)", "Timeline"],
          [
              ["Lean (MVP)",
               "Core 18 features only, single language, basic admin",
               "₹ 8 L  –  ₹ 12 L", "10-14 weeks"],
              ["Mid-Tier (Recommended)",
               "All 31 features, 2 languages, 2 currencies, full admin",
               "₹ 15 L  –  ₹ 22 L", "18-22 weeks"],
              ["Premium (Enterprise)",
               "All 31 + RBAC, multi-vendor, ML recos, native apps",
               "₹ 25 L  –  ₹ 40 L+", "26-36 weeks"],
          ])
    para(doc,
         "Recurring annual operating cost (hosting, SaaS, AMC) "
         "ranges from ₹1.5 L (lean) to ₹6 L+ (premium). Detailed "
         "breakdowns appear in Section 16.")
    h2(doc, "1.2 Why Build, Not Buy?")
    para(doc,
         "Off-the-shelf options (Shopify, WooCommerce, Magento) cover "
         "60-70% of the requested feature set but cannot deliver the "
         "differentiated BOPIS workflow, spin-the-wheel campaign engine, "
         "gift-package builder, and seller-buyer messaging in a single "
         "unified flow without significant custom development on top of "
         "platform fees. A custom Laravel build gives full ownership of "
         "the data, the UX, and the roadmap — and over a three-year "
         "horizon comes out cheaper than Shopify Plus subscriptions plus "
         "third-party apps.")
    pg(doc)

    # ---------- 2. SCOPE ----------
    h1(doc, "2. Project Vision & Scope")
    h2(doc, "2.1 Vision")
    para(doc,
         "ActToAction will operate a customer-first e-commerce storefront "
         "where shoppers can buy online, pick up in store, personalise "
         "gifts, play a spin-the-wheel for offers, message sellers "
         "directly, and pay in their local currency and language. Sellers "
         "and admins will manage catalog, inventory, orders, returns, and "
         "customer conversations from a single unified admin panel.")
    h2(doc, "2.2 In-Scope")
    bullets(doc, [
        "Public storefront (responsive, mobile-first).",
        "Customer account portal with sign-in, wishlist, orders, returns.",
        "Guest checkout path with no forced registration.",
        "Multi-language UI (English + one additional language at launch).",
        "Multi-currency display and checkout.",
        "Full catalog management (products, variants, categories, attributes).",
        "Cart, coupons, gift packages, custom notes.",
        "BOPIS (Buy Online, Pick-up In-Store) workflow.",
        "Payment gateway (Razorpay primary, Stripe optional).",
        "Shipping integration (Delhivery / BlueDart / DHL).",
        "Order management, tracking, invoicing.",
        "Return management workflow.",
        "Push notifications and transactional emails.",
        "Suggestion chatbot for product help.",
        "Spin-the-wheel offer engine.",
        "Admin panel covering every entity above.",
    ])
    h2(doc, "2.3 Out-of-Scope (this proposal)")
    bullets(doc, [
        "Native iOS / Android applications (proposed for Phase 2).",
        "Multi-vendor marketplace (Premium tier add-on).",
        "Subscription / recurring billing.",
        "Loyalty programme with tier-based points (separate proposal).",
        "ERP integration (SAP, Tally, Oracle) — quoted separately.",
        "Migration of legacy customer or order data — quoted separately.",
    ])
    pg(doc)

    # ---------- 3. FEATURE LIST ----------
    h1(doc, "3. Complete Feature List (31 modules)")
    para(doc,
         "Every feature you specified has been catalogued below with a "
         "short positioning note. Detailed specifications, requirements, "
         "and effort estimates appear in the chapters that follow.")
    table(doc,
          ["#", "Feature", "Category", "Tier"],
          [
              ["1", "GO OFFLINE — Get It From Store (BOPIS)", "Fulfilment", "Mid+"],
              ["2", "Multilingual website", "Localisation", "Mid+"],
              ["3", "Currency management", "Localisation", "Mid+"],
              ["4", "Push notifications", "Engagement", "Mid+"],
              ["5", "Sign-in / Sign-up", "Identity", "All"],
              ["6", "Dynamic fields for ads and other content", "Marketing", "Mid+"],
              ["7", "Spin the wheel for offers", "Marketing", "Mid+"],
              ["8", "Payment gateway", "Commerce", "All"],
              ["9", "Merchandise (product catalog)", "Commerce", "All"],
              ["10", "Order management", "Operations", "All"],
              ["11", "Order tracking", "Operations", "All"],
              ["12", "Bill creation / Invoice management", "Operations", "All"],
              ["13", "DHL & other shipping partners", "Logistics", "All"],
              ["14", "Email (transactional + marketing)", "Engagement", "All"],
              ["15", "Return management", "Operations", "All"],
              ["16", "Gift package builder", "Commerce", "Mid+"],
              ["17", "Suggestion chatbot", "Engagement", "Mid+"],
              ["18", "Guest checkout", "Commerce", "All"],
              ["19", "Personalisation request", "Commerce", "Mid+"],
              ["20", "Wishlist & Favourite", "Engagement", "All"],
              ["21", "Cart reminder (abandonment)", "Engagement", "Mid+"],
              ["22", "Product-wise data, state-wise seller", "Analytics", "Mid+"],
              ["23", "Message to buyer (seller-buyer chat)", "Engagement", "Mid+"],
              ["24", "Catalog management", "Operations", "All"],
              ["25", "Coupon listing (Zomato-style chooser)", "Marketing", "Mid+"],
              ["26", "Custom note for personalised product", "Commerce", "Mid+"],
              ["27", "Terms & Conditions page", "Compliance", "All"],
              ["28", "Privacy Policy page", "Compliance", "All"],
              ["29", "Cart", "Commerce", "All"],
              ["30", "Contact us", "Compliance", "All"],
              ["31", "Lazy-load logo & performance tuning", "Performance", "All"],
          ])
    pg(doc)

    # ---------- 4. DETAILED SPECS ----------
    h1(doc, "4. Detailed Feature Specifications")
    para(doc,
         "Each subsection below explains a feature, lists what the user "
         "and admin can do, and notes the key technical components.")

    specs = [
        ("4.1  GO OFFLINE — Get It From Store (BOPIS)",
         "Allows the customer to pay online and pick the product up "
         "from a designated physical store, avoiding shipping fees and "
         "delivery delays.",
         ["Customer: choose 'Pick up in store' at checkout, pick a store from a map/list.",
          "System: hold stock at the chosen store for N hours, send OTP to customer.",
          "Store staff: view pickup queue in store dashboard, verify OTP, mark fulfilled.",
          "Edge cases: store out-of-stock fallback to ship-from-warehouse."]),

        ("4.2  Multilingual Website",
         "Every customer-facing string can be translated into multiple "
         "languages. Launch with English + one regional language; more "
         "can be added later without code changes.",
         ["Translation files under resources/lang/{locale}/.",
          "Auto-detect language from browser, override via URL prefix /hi/.",
          "Right-to-left layout support (for future Arabic/Urdu).",
          "Admin can edit translations live without redeploy."]),

        ("4.3  Currency Management",
         "Prices are stored in a base currency (INR) and displayed in "
         "the customer's chosen currency at live exchange rates.",
         ["FX rates fetched daily from an open API (exchangerate.host).",
          "Customer can override the auto-detected currency.",
          "Checkout charges in INR; display layer converts.",
          "Admin overrides for manual rate adjustments."]),

        ("4.4  Push Notifications",
         "Browser-based and (optionally) mobile push for order updates, "
         "abandoned carts, and promotional campaigns.",
         ["Firebase Cloud Messaging (FCM) or OneSignal SDK.",
          "Admin can compose and segment campaigns by city, gender, last-purchase.",
          "User can opt in/out from account settings."]),

        ("4.5  Sign-in / Sign-up",
         "Multiple authentication options for both convenience and security.",
         ["Email + password (bcrypt-hashed).",
          "OTP login (email or SMS).",
          "Social login (Google, Facebook) optional.",
          "Forgot-password flow with rate-limited reset links."]),

        ("4.6  Dynamic Fields for Ads and Other Content",
         "Admin can create custom fields on any entity (product, banner, "
         "promo) without changing code, enabling rapid campaign launches.",
         ["JSON column on each major entity for ad-hoc attributes.",
          "Admin UI to define field name, type, and validation.",
          "Used by ad slots, banners, and landing-page blocks."]),

        ("4.7  Spin the Wheel for Offers",
         "A gamified promotional widget that customers can spin once per "
         "session for a discount coupon or a free product sample.",
         ["Configurable wheel segments (offer, probability, coupon code).",
          "Anti-fraud: one spin per IP + email per 24h.",
          "Spins logged for analytics; coupon usage tracked.",
          "Mobile-friendly canvas animation."]),

        ("4.8  Payment Gateway",
         "Inline checkout integration with Razorpay (primary) and Stripe "
         "(optional for international cards).",
         ["Razorpay Standard Checkout (cards, UPI, netbanking, wallets, EMI).",
          "Server-side HMAC signature verification.",
          "Order-payment-reconciliation with unique idempotency keys.",
          "Webhook handlers for refunds and disputes."]),

        ("4.9  Merchandise (Product Catalog)",
         "Full product catalog with variants, attributes, images, SEO "
         "metadata, and inventory.",
         ["Products → Variants → SKU, with size/colour/material axes.",
          "Multiple images per product, lazy-loaded with WebP.",
          "Per-variant inventory and pricing.",
          "SEO slug, meta title, meta description, OG image."]),

        ("4.10  Order Management",
         "Admin views every order with filters by date, status, customer, "
         "payment method, and channel.",
         ["States: pending → paid → packed → shipped → delivered → closed.",
          "Bulk actions: mark shipped, print labels, download invoices.",
          "Notes per order for internal use.",
          "CSV export for accounting reconciliation."]),

        ("4.11  Order Tracking",
         "Customer-facing real-time tracking with timeline and courier "
         "status.",
         ["Polls courier APIs every 4h, persists status events.",
          "Customer page shows progress bar + ETA.",
          "Push + email notification on each milestone."]),

        ("4.12  Bill Creation / Invoice Management",
         "Auto-generated GST-compliant invoices in PDF; downloadable by "
         "customer and admin.",
         ["barryvdh/laravel-dompdf for PDF generation.",
          "Sequential invoice numbering with financial-year reset.",
          "GSTIN, HSN/SAC codes, line-item tax breakup.",
          "Credit notes for refunds and returns."]),

        ("4.13  DHL & Other Shipping Partners",
         "Programmatic integration with major couriers to create "
         "shipments, fetch tracking, and print labels.",
         ["Adapters for DHL, BlueDart, Delhivery, Shiprocket aggregator.",
          "Pick best rate per pincode + weight.",
          "Auto-AWB generation on order pack.",
          "Cancellation and reschedule support."]),

        ("4.14  Email (Transactional + Marketing)",
         "Templated emails for every key event, plus marketing campaigns.",
         ["SMTP via SendGrid / Amazon SES.",
          "Templates with merge variables in admin.",
          "email_logs for every sent message, retryable on failure.",
          "Unsubscribe link mandatory on marketing emails."]),

        ("4.15  Return Management",
         "Customer-initiated returns with reason codes, photo uploads, "
         "and admin approval workflow.",
         ["Return window configurable per product category.",
          "States: requested → approved → picked up → received → refunded.",
          "Reverse logistics with courier API.",
          "Refund to original payment method via Razorpay."]),

        ("4.16  Gift Package Builder",
         "Customer selects multiple products, adds a custom card, and "
         "buys as a single gift bundle.",
         ["Drag-and-drop builder with live total.",
          "Optional gift wrap (paid add-on).",
          "Greeting card with personal message.",
          "Direct ship to recipient address."]),

        ("4.17  Suggestion Chatbot",
         "An on-site assistant that answers FAQs and recommends products "
         "based on browsing behaviour.",
         ["FAQ-first widget; falls back to support ticket if unanswered.",
          "Hooks into product catalog to suggest similar items.",
          "Optionally backed by an LLM (Claude / GPT) — extra SaaS cost.",
          "Captures unanswered questions for admin curation."]),

        ("4.18  Guest Checkout",
         "Allow purchase without forced registration.",
         ["Email captured at checkout; account auto-created post-purchase optionally.",
          "Guest order lookup by order-id + email.",
          "Reduces checkout abandonment significantly."]),

        ("4.19  Personalisation Request",
         "Customer can request a personalised product variant (engraving, "
         "custom colour, custom size) that the seller fulfils manually.",
         ["Dedicated form on product page with text + file upload.",
          "Routed to seller dashboard for quote/confirmation.",
          "Custom pricing applied per request."]),

        ("4.20  Wishlist & Favourites",
         "Customers save products for later viewing.",
         ["Logged-in users: persisted to DB.",
          "Guest users: stored in localStorage, merged on login.",
          "Wishlist back-in-stock and price-drop alerts."]),

        ("4.21  Cart Reminder (Abandonment)",
         "Email and push reminders for carts left without checkout.",
         ["Trigger after 1h, 24h, 72h windows.",
          "Optional coupon attached to incentivise return.",
          "Stop reminders once converted or unsubscribed."]),

        ("4.22  Product-wise Data, State-wise Seller",
         "Analytics that pivot products against geography and seller "
         "performance — which products sell best in which state, which "
         "seller leads each region.",
         ["Pre-aggregated dashboards (daily cron).",
          "Drilldowns by state, district, pincode.",
          "Seller leaderboard with rank by units and revenue.",
          "Exportable PDF / Excel reports for management review."]),

        ("4.23  Message to Buyer (Seller-Buyer Chat)",
         "Real-time messaging between sellers and buyers for "
         "clarifications and personalised quotes.",
         ["Per-order conversation thread (audit trail).",
          "Image attachments supported.",
          "Push + email notification on new message.",
          "Admin oversight: ability to read any thread for moderation."]),

        ("4.24  Catalog Management",
         "Comprehensive admin tools for products, categories, attributes, "
         "stock, and pricing.",
         ["Bulk CSV import / export.",
          "Drag-and-drop category reordering.",
          "Inventory thresholds with low-stock alerts.",
          "Per-channel availability (web vs in-store)."]),

        ("4.25  Coupon Code Listing (Zomato-style)",
         "Customers see all eligible coupons at checkout, click to apply.",
         ["Coupon engine supports flat, percentage, BOGO, free shipping, free gift.",
          "Filters by category, brand, first-time, minimum cart value.",
          "Showcased on coupon listing page for browse-and-grab UX.",
          "Per-customer redemption limit + global cap."]),

        ("4.26  Custom Note for Personalised Product",
         "Free-text field on cart for instructions to seller (e.g. "
         "'Please pack in red ribbon').",
         ["Visible in seller dashboard alongside the order.",
          "Optional photo / file upload.",
          "Printable on packing slip."]),

        ("4.27  Terms & Conditions Page",
         "A statutory page outlining the contract between ActToAction and "
         "the customer.",
         ["Admin-editable Markdown / rich text.",
          "Versioned: customer acceptance captured at registration.",
          "Linked from footer and checkout."]),

        ("4.28  Privacy Policy Page",
         "GDPR / DPDP-compliant policy describing data collection and use.",
         ["Admin-editable.",
          "DPDP-2023 (India) cookie banner with consent log.",
          "Linked from footer, checkout, and sign-up."]),

        ("4.29  Cart",
         "The standard cart page with edit-quantity, remove, apply-coupon, "
         "and proceed-to-checkout actions.",
         ["Sticky 'Proceed to Checkout' button on mobile.",
          "Free-shipping and offer eligibility banners.",
          "Save-for-later moves item to wishlist."]),

        ("4.30  Contact Us",
         "Contact form, address, phone, email, business hours, and a map.",
         ["Form submissions persisted and emailed to support.",
          "Google Maps embed for the store address.",
          "Rate-limited to deter spam; honeypot field."]),

        ("4.31  Lazy-Load Logo & Performance Tuning",
         "All images lazy-loaded; logo uses inline SVG for instant "
         "render. Sitewide perf budget enforced.",
         ["loading='lazy' on every non-fold image.",
          "Inline critical CSS; defer non-critical JS.",
          "WebP/AVIF where browser supports.",
          "Lighthouse target ≥ 90 mobile, ≥ 95 desktop."]),
    ]
    for title, intro, items in specs:
        h2(doc, title)
        para(doc, intro)
        bullets(doc, items)
    pg(doc)

    # ---------- 5. FRs ----------
    h1(doc, "5. Functional Requirements")
    para(doc,
         "Functional requirements are organised by module. The list is "
         "the minimum acceptance bar; additional polish and edge-cases "
         "are part of the standard delivery quality.")
    fr_blocks = [
        ("5.1  Identity",
         ["FR-1: System shall allow sign-up with email + password.",
          "FR-2: System shall allow OTP-based sign-in.",
          "FR-3: System shall support forgot-password with time-limited token.",
          "FR-4: System shall lock the account after 5 failed attempts in 5 minutes.",
          "FR-5: System shall allow guest checkout without registration."]),
        ("5.2  Catalog",
         ["FR-6: Admin shall create/edit/delete products and variants.",
          "FR-7: Admin shall upload up to 8 images per product.",
          "FR-8: Customer shall search products by name, SKU, brand.",
          "FR-9: Customer shall filter by category, price, attribute, rating.",
          "FR-10: System shall hide out-of-stock variants from selection by default."]),
        ("5.3  Cart & Checkout",
         ["FR-11: Customer shall add/remove items, update quantity in cart.",
          "FR-12: Customer shall apply at most one coupon at checkout.",
          "FR-13: System shall present available coupons on a dedicated chooser.",
          "FR-14: Customer shall pay via Razorpay; verification done server-side.",
          "FR-15: Customer shall optionally request store pickup (BOPIS)."]),
        ("5.4  Orders & Returns",
         ["FR-16: Customer shall view order history and track open orders.",
          "FR-17: Customer shall initiate return within configurable window.",
          "FR-18: System shall trigger reverse pickup with courier.",
          "FR-19: System shall refund to original payment method on receipt of returned goods.",
          "FR-20: Admin shall view aging reports of unfulfilled / pending-pickup orders."]),
        ("5.5  Communication",
         ["FR-21: System shall send transactional emails for sign-up, order, shipment, refund.",
          "FR-22: System shall send web push notifications for the same events.",
          "FR-23: Customer shall message seller from the order details page.",
          "FR-24: System shall trigger cart-abandonment email after 1h, 24h, 72h."]),
        ("5.6  Marketing",
         ["FR-25: Admin shall configure spin-the-wheel segments and probabilities.",
          "FR-26: Admin shall create coupon codes (flat, %, BOGO, free shipping).",
          "FR-27: Admin shall schedule push and email campaigns.",
          "FR-28: System shall show personalisation request form on selected SKUs."]),
        ("5.7  Operations",
         ["FR-29: Admin shall view orders pivoted by state, city, seller.",
          "FR-30: Admin shall download GST-compliant invoices.",
          "FR-31: Admin shall manage stock at multiple warehouses + stores.",
          "FR-32: Admin shall edit Terms, Privacy, Contact pages."]),
    ]
    for title, items in fr_blocks:
        h2(doc, title)
        bullets(doc, items)
    pg(doc)

    # ---------- 6. NFRs ----------
    h1(doc, "6. Non-Functional Requirements")
    table(doc,
          ["Category", "Requirement"],
          [
              ["Performance", "P95 page load ≤ 2.0s on 4G; Lighthouse mobile ≥ 90."],
              ["Scalability", "Support 500 concurrent checkouts without queueing."],
              ["Availability", "99.5% monthly uptime SLA excluding scheduled windows."],
              ["Security", "HTTPS everywhere, OWASP Top 10 covered, signed payment callbacks."],
              ["Compliance", "DPDP-2023 (India), GST e-invoicing, PCI-DSS scope minimisation via Razorpay."],
              ["Accessibility", "WCAG 2.1 AA for public storefront."],
              ["Browser support", "Chrome, Firefox, Safari, Edge — current + 1 prior major."],
              ["Mobile", "Responsive from 320px upwards."],
              ["i18n", "English + 1 regional language at launch; add more without code change."],
              ["Backups", "Nightly DB + media snapshots, 14-day rolling retention."],
              ["Observability", "Application logs, error tracker (Sentry), uptime monitor."],
              ["Maintainability", "PSR-12 PHP, ESLint JS, conventional commits, CI on push."],
          ])
    pg(doc)

    # ---------- 7. TECH STACK ----------
    h1(doc, "7. Technology Stack")
    table(doc,
          ["Layer", "Recommended", "Why"],
          [
              ["Backend framework", "Laravel 12 (PHP 8.2+)", "Battle-tested, deep talent pool, fast to build."],
              ["Database", "MySQL 8.0", "Stable, well-supported, ACID."],
              ["Cache / Queue", "Redis 7", "Sessions, queues, abandoned-cart timers."],
              ["Search", "MeiliSearch / Elasticsearch", "Sub-second product search."],
              ["Frontend", "Blade + Alpine.js + Tailwind", "Server-rendered for SEO, lightweight JS."],
              ["Mobile-friendly", "Responsive Tailwind", "Single codebase, all devices."],
              ["Payments", "Razorpay (primary), Stripe (intl)", "Local methods + global cards."],
              ["Shipping", "Shiprocket aggregator + direct APIs", "BlueDart, Delhivery, DHL coverage."],
              ["Push", "OneSignal / Firebase FCM", "Web + future mobile parity."],
              ["Email", "SendGrid / Amazon SES", "High deliverability."],
              ["Chatbot", "Rule engine + optional Claude/GPT", "Cheap baseline, smart upgrade."],
              ["PDF", "barryvdh/laravel-dompdf", "Invoices and receipts."],
              ["Hosting", "AWS Mumbai / DigitalOcean BLR", "Low latency to Indian customers."],
              ["CDN", "Cloudflare", "WAF + edge cache included."],
              ["Monitoring", "Sentry + UptimeRobot + LogTail", "Visibility on errors and SLAs."],
              ["CI / CD", "GitHub Actions", "Free for moderate use, simple YAML."],
          ])
    pg(doc)

    # ---------- 8. ARCH ----------
    h1(doc, "8. Proposed System Architecture")
    img(doc, arch, "Figure 8.1 — High-level architecture spanning users, edge, application, services, data, and integrations.", width=6.4)
    para(doc,
         "Traffic terminates at Cloudflare (CDN + WAF), then hits the "
         "Laravel application behind Nginx. The application is split "
         "logically into cohesive service modules (catalog, cart, "
         "orders, payments, returns, BOPIS, coupons, i18n, chatbot, "
         "push, email, gift builder, analytics, spin-wheel). All write "
         "to MySQL (primary store) and use Redis for cache, sessions, "
         "and queued background jobs. Static media and customer uploads "
         "go to S3-compatible storage. Razorpay, courier APIs, push "
         "providers, and SMTP relays are accessed over HTTPS.")
    pg(doc)

    # ---------- 9. USER FLOW ----------
    h1(doc, "9. Customer Purchase Flow")
    img(doc, flow, "Figure 9.1 — Customer journey from landing to fulfilment, with optional return.", width=6.4)
    para(doc,
         "The flow captures the dominant happy path. Side paths "
         "(abandon cart, OTP failure, payment failure, partial fulfilment, "
         "split shipment) are handled by dedicated controllers and "
         "covered in detailed sequence diagrams during the design phase.")
    pg(doc)

    # ---------- 10. DB MODULES ----------
    h1(doc, "10. Database Modules")
    para(doc,
         "The schema will span approximately 55-70 tables grouped by "
         "functional cluster. Indicative tables per cluster are listed "
         "below; detailed ERDs will be produced during the design phase.")
    table(doc,
          ["Cluster", "Indicative Tables"],
          [
              ["Identity", "users, roles, otps, sessions, password_resets, addresses"],
              ["Catalog", "categories, products, product_variants, attributes, attribute_values, product_images, brands, tags"],
              ["Inventory", "warehouses, stores, stock_levels, stock_movements"],
              ["Pricing", "currencies, fx_rates, price_lists, taxes"],
              ["Cart & Checkout", "carts, cart_items, wishlist, saved_for_later, coupons, coupon_usages"],
              ["Orders", "orders, order_items, order_status_history, payment_attempts, payments, refunds"],
              ["Fulfilment", "shipments, shipment_events, pickup_slots, bopis_holds"],
              ["Returns", "return_requests, return_items, return_pickups, credit_notes"],
              ["Marketing", "campaigns, spin_wheels, spin_segments, spin_attempts, banners, ads"],
              ["Engagement", "push_subscriptions, push_messages, email_logs, email_templates, abandonment_jobs"],
              ["Messaging", "conversations, messages, message_attachments"],
              ["Content", "pages (T&C, Privacy, Contact), faqs, blog_posts"],
              ["Localisation", "translations, locales"],
              ["Audit", "activity_logs"],
          ])
    pg(doc)

    # ---------- 11. INTEGRATIONS ----------
    h1(doc, "11. Third-Party Integrations & SaaS Costs")
    table(doc,
          ["Service", "Purpose", "Indicative Cost (INR/mo or fee)"],
          [
              ["Razorpay", "Payment gateway",
               "2% per transaction (domestic), 3% intl"],
              ["Shiprocket", "Courier aggregation",
               "Pay-per-shipment, ~₹30-80 per AWB"],
              ["DHL / BlueDart / Delhivery", "Direct courier APIs",
               "Per-shipment, negotiated rates"],
              ["SendGrid", "Transactional email",
               "Free tier 100/day; ~₹1,500/mo for 50k"],
              ["OneSignal / FCM", "Push notifications",
               "Free for basic; ~₹6,000+/mo for advanced"],
              ["Cloudflare", "CDN + WAF",
               "Free → ~₹1,700/mo (Pro plan)"],
              ["AWS Mumbai", "Hosting",
               "~₹6,000-25,000/mo depending on traffic"],
              ["Sentry", "Error tracking",
               "Free tier; ~₹2,200/mo Team plan"],
              ["UptimeRobot", "Uptime monitoring", "Free for 50 monitors"],
              ["LogTail / Datadog", "Log management",
               "~₹2,000-8,000/mo"],
              ["Google Maps", "Address autocomplete",
               "Free tier covers ~28k loads/mo"],
              ["Claude / GPT (optional)", "Chatbot LLM",
               "Pay-per-token, ~₹4,000-15,000/mo at low scale"],
              ["S3 (or DO Spaces)", "Object storage",
               "~₹500-2,500/mo for first TB"],
              ["Domain + TLS", "Domain & SSL",
               "~₹1,500/yr (Let's Encrypt TLS is free)"],
          ])
    pg(doc)

    # ---------- 12. GANTT ----------
    h1(doc, "12. Project Schedule — 22 Weeks")
    img(doc, gantt, "Figure 12.1 — Gantt chart of the 22-week mid-tier build schedule.", width=6.5)
    h2(doc, "12.1 Phase Summary")
    table(doc,
          ["Phase", "Weeks", "Key Deliverables"],
          [
              ["Discovery & Wireframes", "1-3", "Wireframes, sitemap, user-journey map"],
              ["UI/UX Design", "3-6", "High-fidelity mockups, design system"],
              ["Backend Foundation", "5-7", "Auth, base schema, deploy pipeline"],
              ["Catalog + Cart", "7-10", "Products, variants, search, cart"],
              ["Payments + Orders", "10-12", "Razorpay, order workflow"],
              ["i18n + Multi-currency", "10-11", "Translations and FX"],
              ["Shipping (DHL etc.)", "12-13", "Courier API adapters"],
              ["BOPIS", "13-14", "Store pickup flow"],
              ["Returns + Invoicing", "13-14", "Return workflow, GST PDF"],
              ["Coupons + Spin Wheel", "14-15", "Promotion engine"],
              ["Gift Builder + Notes", "15-16", "Bundle UI, custom notes"],
              ["Chatbot + Push + Email", "15-17", "Engagement layer"],
              ["Admin Panel", "5-16", "All admin CRUDs (runs in parallel)"],
              ["QA + UAT", "18-20", "Test plan execution, bug fixes"],
              ["Deployment + Handover", "21-22", "Production go-live, training"],
          ])
    pg(doc)

    # ---------- 13. EFFORT ----------
    h1(doc, "13. Per-Feature Effort Estimate")
    img(doc, feff, "Figure 13.1 — Effort per feature at a blended rate of ₹900/hr.", width=6.4)
    para(doc,
         "These numbers represent the development effort only. Each "
         "feature additionally consumes design, QA, project management, "
         "and deployment time which is rolled up under the central cost "
         "breakdown in Section 14.")
    pg(doc)

    # ---------- 14. COST BREAKDOWN ----------
    h1(doc, "14. India-Market Cost Breakdown")
    para(doc,
         "All figures are in Indian Rupees and exclude GST (18%). Rates "
         "assume a blended team of senior + mid-level Indian developers "
         "at ₹800-1,200 per hour.")
    h2(doc, "14.1 Mid-Tier Detailed Breakdown (Recommended)")
    table(doc,
          ["Line Item", "Effort", "Rate", "Cost"],
          [
              ["UI/UX Design (wireframes + hi-fi mockups)", "150h", "₹1,000", "₹ 1,50,000"],
              ["Backend Development (all modules)", "750h", "₹800", "₹ 6,00,000"],
              ["Frontend Development (Blade + Alpine + Tailwind)", "450h", "₹900", "₹ 4,00,000"],
              ["Third-party Integrations (Razorpay, DHL, push, email)", "280h", "₹900", "₹ 2,50,000"],
              ["Admin Panel (40+ CRUD screens)", "240h", "₹850", "₹ 2,00,000"],
              ["QA & Testing (test plan + execution)", "180h", "₹850", "₹ 1,50,000"],
              ["Project Management & Coordination", "150h", "₹800", "₹ 1,20,000"],
              ["DevOps & Deployment", "60h", "₹850", "₹ 50,000"],
              ["Contingency (~10%)", "—", "—", "₹ 1,80,000"],
              ["Subtotal (Mid-Tier)", "", "", "₹ 21,00,000"],
              ["GST @ 18%", "", "", "₹ 3,78,000"],
              ["Grand Total (incl. GST)", "", "", "₹ 24,78,000"],
          ])
    img(doc, pie, "Figure 14.1 — Mid-tier cost breakdown by work category.", width=6.0)
    pg(doc)

    # ---------- 15. TIERS ----------
    h1(doc, "15. Three Engagement Tiers")
    img(doc, tiers, "Figure 15.1 — Low / high cost band per tier.", width=6.3)
    h2(doc, "15.1 Lean (MVP) — ₹ 8 L to ₹ 12 L")
    para(doc,
         "Aimed at validating the business model with the minimum viable "
         "set of features.")
    bullets(doc, [
        "Catalog, cart, guest checkout, coupons.",
        "Razorpay payment, basic order management.",
        "Single language, single currency (INR).",
        "Email transactional only (no push, no chatbot).",
        "Shiprocket aggregator only (no direct courier APIs).",
        "Basic admin panel.",
        "10-14 weeks timeline.",
    ])
    h2(doc, "15.2 Mid-Tier (Recommended) — ₹ 15 L to ₹ 22 L")
    para(doc,
         "The full feature set you requested, delivered in 18-22 weeks.")
    bullets(doc, [
        "All 31 features from Section 3.",
        "Two languages, two currencies.",
        "Full admin panel with role-aware actions.",
        "Push, chatbot, gift builder, spin-the-wheel.",
        "BOPIS, return workflow, GST invoices.",
        "Direct DHL / BlueDart / Delhivery integration.",
    ])
    h2(doc, "15.3 Premium (Enterprise) — ₹ 25 L to ₹ 40 L+")
    para(doc,
         "Adds enterprise capabilities on top of the mid-tier scope.")
    bullets(doc, [
        "Native iOS + Android apps (React Native or Flutter).",
        "Multi-vendor marketplace with seller onboarding.",
        "ML-based product recommendations.",
        "Advanced analytics dashboards (cohorts, RFM, LTV).",
        "Role-based access control with audit trail.",
        "ERP / accounting integration (Tally, Zoho Books).",
        "Loyalty programme.",
        "Dedicated SRE + 24×7 monitoring.",
    ])
    pg(doc)

    # ---------- 16. RECURRING ----------
    h1(doc, "16. Recurring Annual Costs")
    table(doc,
          ["Item", "Lean (₹/yr)", "Mid-Tier (₹/yr)", "Premium (₹/yr)"],
          [
              ["Hosting (AWS / DO)", "₹ 50,000", "₹ 1,50,000", "₹ 4,00,000"],
              ["Domain + TLS", "₹ 1,500", "₹ 1,500", "₹ 1,500"],
              ["CDN (Cloudflare)", "Free", "₹ 20,000", "₹ 60,000"],
              ["Email (SendGrid)", "Free", "₹ 18,000", "₹ 60,000"],
              ["Push (OneSignal)", "Free", "₹ 18,000", "₹ 72,000"],
              ["Error tracking (Sentry)", "Free", "₹ 26,000", "₹ 60,000"],
              ["Backups (S3 etc.)", "₹ 6,000", "₹ 18,000", "₹ 60,000"],
              ["LLM (chatbot) optional", "—", "₹ 60,000", "₹ 1,80,000"],
              ["Logs (LogTail)", "—", "₹ 24,000", "₹ 96,000"],
              ["AMC (this vendor)", "₹ 60,000", "₹ 2,40,000", "₹ 6,00,000"],
              ["Approximate Total", "₹ 1.2 L", "₹ 5.7 L", "₹ 15 L"],
          ])
    para(doc,
         "Razorpay and courier costs are transaction-volume-driven and "
         "billed as a percentage of GMV — not included above. Expect "
         "Razorpay at ~2% of revenue and courier at ₹40-100 per order.")
    pg(doc)

    # ---------- 17. MILESTONES ----------
    h1(doc, "17. Payment Milestones")
    para(doc,
         "We propose milestone-based billing tied to deliverables, "
         "rather than time-based invoicing. This aligns risk and "
         "reward.")
    table(doc,
          ["Milestone", "% of Total", "Trigger"],
          [
              ["M1: Project Kickoff", "20%", "Contract signed, kickoff meeting completed."],
              ["M2: Design Sign-off", "15%", "All hi-fi mockups approved."],
              ["M3: Catalog + Cart UAT", "20%", "Customer can browse + add to cart + sign in."],
              ["M4: Payments + Orders UAT", "20%", "Customer can place a paid order; admin can view it."],
              ["M5: Feature Completion", "15%", "All 31 features implemented and demoed."],
              ["M6: Go-Live", "10%", "Production deployment, training delivered."],
          ])
    pg(doc)

    # ---------- 18. ASSUMPTIONS ----------
    h1(doc, "18. Assumptions, Inclusions & Exclusions")
    h2(doc, "18.1 Inclusions")
    bullets(doc, [
        "Source code with full ownership transferred on final payment.",
        "30 days of post-launch bug-fix support included.",
        "Two rounds of design revisions per screen.",
        "Two rounds of stakeholder feedback per development sprint.",
        "Knowledge transfer session + admin user manual.",
        "Basic SEO setup (sitemap, robots.txt, OG tags, structured data).",
    ])
    h2(doc, "18.2 Exclusions")
    bullets(doc, [
        "Content writing for product descriptions, blogs, T&C, Privacy.",
        "Photography of products.",
        "Translation of UI strings beyond the first regional language (English + 1).",
        "Logo and brand identity design (assumed already in place).",
        "Marketing campaigns post-launch.",
        "Payment gateway and courier merchant account setup with the respective vendors.",
        "Server / hosting fees during the project (we bill, you pay vendor directly post go-live).",
    ])
    h2(doc, "18.3 Assumptions")
    bullets(doc, [
        "Client provides timely feedback (≤ 2 business days per ask).",
        "Client provides production content and assets by week 14.",
        "Razorpay merchant account is active by week 8.",
        "Courier accounts (DHL / BlueDart / Delhivery) are open by week 11.",
        "No more than 31 features within the agreed scope; additions billed per change-request.",
    ])
    pg(doc)

    # ---------- 19. RISKS ----------
    h1(doc, "19. Risks & Mitigation")
    table(doc,
          ["Risk", "Likelihood", "Impact", "Mitigation"],
          [
              ["Scope creep beyond 31 features", "High", "High",
               "Strict change-request process with re-estimates."],
              ["Courier API rate-limits and instability", "Medium", "Medium",
               "Aggregator (Shiprocket) as fallback path."],
              ["Razorpay onboarding delays", "Medium", "High",
               "Start onboarding day 1; test mode dev until live."],
              ["Multilingual translation quality", "Medium", "Medium",
               "Use professional translator, not machine-only."],
              ["BOPIS adoption lower than expected", "Medium", "Low",
               "Build as a feature flag; can be hidden if unused."],
              ["Payment fraud", "Low", "High",
               "Razorpay risk engine + signature verification + velocity rules."],
              ["DPDP-2023 enforcement", "Medium", "High",
               "Consent log, opt-out flow, data export endpoint."],
              ["Performance regressions over time", "Medium", "Medium",
               "Lighthouse CI + weekly perf review."],
          ])
    pg(doc)

    # ---------- 20. AMC ----------
    h1(doc, "20. Post-Launch Support (AMC)")
    para(doc,
         "Annual Maintenance Contract options are available after the "
         "30-day free support window expires.")
    table(doc,
          ["AMC Plan", "Inclusions", "Annual Fee"],
          [
              ["Bronze",
               "Email support, bug fixes only, monthly patching, business hours response (24h SLA)",
               "₹ 60,000"],
              ["Silver",
               "Bronze + minor enhancements (up to 20h/mo), 12h SLA, monthly performance review",
               "₹ 2,40,000"],
              ["Gold",
               "Silver + dedicated PM, 4h SLA, quarterly roadmap planning, 40h/mo enhancements",
               "₹ 6,00,000"],
              ["Platinum",
               "Gold + 24×7 on-call, 1h SLA on critical, monthly architecture review",
               "₹ 12,00,000+"],
          ])
    pg(doc)

    # ---------- 21. ACCEPTANCE ----------
    h1(doc, "21. Acceptance & Sign-Off")
    para(doc,
         "By signing below, the parties agree to the scope, milestones, "
         "and commercial terms set out in this document. Any deviation "
         "shall be governed by the change-request process described in "
         "Section 18.")
    para(doc, "")
    para(doc, "")
    t = doc.add_table(rows=4, cols=2)
    cells = [
        ("For ActToAction:", "For the Development Vendor:"),
        ("Name: __________________________", "Name: __________________________"),
        ("Designation: ____________________", "Designation: ____________________"),
        ("Signature: ______________________", "Signature: ______________________"),
    ]
    for i, (left, right) in enumerate(cells):
        for j, txt in enumerate((left, right)):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_font(r, size=12, bold=(i == 0))
            set_p(p, sa=10)
    para(doc, "")
    para(doc, "Date: __________________________")
    para(doc,
         "This proposal is valid for 30 days from the version date on "
         "the cover page.",
         indent=0, justify=False)

    out = os.path.join(OUTDIR, "ActToAction_Proposal_and_Cost.docx")
    doc.save(out)
    print(f"Saved to {out}")
    return out


if __name__ == "__main__":
    build()
